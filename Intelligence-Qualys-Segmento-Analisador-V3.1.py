import customtkinter as ctk
from tkinter import messagebox, filedialog
import requests
import time
import os
import xml.etree.ElementTree as ET 
import csv
import re
from io import StringIO
from dotenv import load_dotenv
from datetime import datetime, timezone, timedelta
import pandas as pd
import google.generativeai as genai
from groq import Groq
import sys
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.gridspec import GridSpec
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import threading
import traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
import random
from dataclasses import dataclass
from typing import Dict, List, Optional
import glob
import webbrowser
from itertools import cycle
from flask import Flask, render_template_string, jsonify
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import sqlite3

# --- CONFIGURAÇÃO INICIAL ---
load_dotenv()
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

# Configuração Matplotlib para Threads (Sem GUI para evitar crash)
matplotlib.use('Agg') 

GEMINI_KEY = os.getenv("GEMINI_API_KEY")
GROQ_KEY = os.getenv("GROQ_API_KEY")

# --- LÓGICA DE CARREGAMENTO DE MÚLTIPLAS CHAVES NVD ---
nvd_keys_env = os.getenv("NVD_API_KEYS")
if nvd_keys_env:
    NVD_API_KEYS = [k.strip() for k in nvd_keys_env.split(',') if k.strip()]
else:
    single_key = os.getenv("NVD_API_KEY")
    NVD_API_KEYS = [single_key] if single_key else []

AI_AVAILABLE = False
if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)
    AI_AVAILABLE = True

# --- GLOBAL SHARED DATA (Ponte entre GUI e Flask) ---
class SharedData:
    dataframe = pd.DataFrame()
    client_name = "Nenhum Cliente Carregado"
    updated_at = datetime.now()

shared_data = SharedData()

# --- INTEGRAÇÃO NVD / THREAT INTEL (COM ROUND ROBIN E SMART TTL) ---

class PersistentNVDCache:
    def __init__(self, db_path="nvd_cache.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Inicializa o banco de dados SQLite para cache."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS nvd_cves (
                    cve_id TEXT PRIMARY KEY,
                    data TEXT,
                    updated_at TEXT
                )
            """)

    def save_cve(self, cve_id: str, data: dict):
        """Salva ou atualiza um CVE no cache."""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    "INSERT OR REPLACE INTO nvd_cves (cve_id, data, updated_at) VALUES (?, ?, ?)",
                    (cve_id, json.dumps(data), datetime.now().isoformat())
                )
        except Exception as e:
            print(f"Erro ao salvar cache SQLite: {e}")

    def get_cve(self, cve_id):
        """Recupera um CVE usando a estratégia Smart TTL."""
        try:
            conn = sqlite3.connect(self.db_path, check_same_thread=False)
            cursor = conn.cursor()
            cursor.execute("SELECT data, updated_at FROM nvd_cves WHERE cve_id = ?", (cve_id,))
            row = cursor.fetchone()
            conn.close()

            if row:
                data_json, updated_at_str = row
                data_dict = json.loads(data_json)
                updated_at = datetime.fromisoformat(updated_at_str)
                age = datetime.now() - updated_at

                # --- LÓGICA INTELIGENTE (SMART TTL) ---
                vuln_status = data_dict.get('vulnStatus', '') 
                
                # Se ainda não foi analisada, validade curta (1 dia)
                if 'AWAITING' in vuln_status.upper() or 'UNDERGOING' in vuln_status.upper():
                    limit = timedelta(days=1)
                else:
                    # Se já foi analisada ou rejeitada, validade longa (5 dias)
                    limit = timedelta(days=5) 

                if age < limit:
                    return data_dict
                
                return None
            return None
        except Exception as e:
            return None

class PersistentQualysKBCache:
    def __init__(self, db_path="qualys_kb_cache.db"):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        """Inicializa o banco de dados SQLite para cache do KB Qualys."""
        with sqlite3.connect(self.db_path) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS qualys_kb (
                    qid TEXT PRIMARY KEY,
                    data TEXT,
                    updated_at TEXT
                )
            """)

    def save_qid(self, qid: str, data: dict):
        """Salva ou atualiza um QID no cache."""
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute(
                    "INSERT OR REPLACE INTO qualys_kb (qid, data, updated_at) VALUES (?, ?, ?)",
                    (str(qid), json.dumps(data), datetime.now().isoformat())
                )
        except Exception as e:
            print(f"Erro ao salvar cache Qualys KB: {e}")

    def get_qid(self, qid):
        """Recupera dados do QID se existirem (Validade longa: 30 dias)."""
        try:
            conn = sqlite3.connect(self.db_path, check_same_thread=False)
            cursor = conn.cursor()
            cursor.execute("SELECT data, updated_at FROM qualys_kb WHERE qid = ?", (str(qid),))
            row = cursor.fetchone()
            conn.close()

            if row:
                data_json, updated_at_str = row
                updated_at = datetime.fromisoformat(updated_at_str)
                if (datetime.now() - updated_at).days < 30:
                    return json.loads(data_json)
            return None
        except Exception as e:
            return None

class NVDIntegration:
    BASE_URL = "https://services.nvd.nist.gov/rest/json/cves/2.0"

    def __init__(self, api_keys: List[str]):
        self.api_keys = api_keys
        self.lock = threading.Lock()
        self.cache = PersistentNVDCache()
        
        if not self.api_keys:
            self.key_pool = [{'key': None, 'history': [], 'limit': 5}]
        else:
            self.key_pool = [{'key': k, 'history': [], 'limit': 50} for k in self.api_keys]
            
        self.window_seconds = 30
        self.pool_cycle = cycle(self.key_pool)

    def _get_valid_header_and_wait(self):
        while True:
            with self.lock:
                now = time.time()
                best_sleep_time = float('inf')
                for _ in range(len(self.key_pool)):
                    key_obj = next(self.pool_cycle)
                    key_obj['history'] = [t for t in key_obj['history'] if now - t < self.window_seconds]
                    if len(key_obj['history']) < key_obj['limit']:
                        key_obj['history'].append(now)
                        headers = {}
                        if key_obj['key']: headers['apiKey'] = key_obj['key']
                        return headers
                    if key_obj['history']:
                        wait = self.window_seconds - (now - key_obj['history'][0])
                        if wait < best_sleep_time: best_sleep_time = wait
            sleep_needed = max(0.1, min(best_sleep_time, 1.0)) + random.uniform(0.1, 0.5)
            time.sleep(sleep_needed)

    def get_cve_details(self, cve_id: str) -> Optional[Dict]:
        cached_data = self.cache.get_cve(cve_id)
        if cached_data: return cached_data

        try:
            headers = self._get_valid_header_and_wait()
            response = requests.get(self.BASE_URL, params={'cveId': cve_id}, headers=headers, timeout=30)
            if response.status_code == 200:
                data = response.json()
                if data.get('vulnerabilities'):
                    parsed_data = self._parse_cve(data['vulnerabilities'][0])
                    self.cache.save_cve(cve_id, parsed_data)
                    return parsed_data
            elif response.status_code in [429, 503]:
                time.sleep(2); return self.get_cve_details(cve_id)
            return None
        except Exception as e: return None

    def _parse_cve(self, cve_data: dict) -> dict:
        cve = cve_data.get('cve', {})
        metrics = cve.get('metrics', {})
        cvss_v31 = metrics.get('cvssMetricV31', [{}])[0].get('cvssData', {}) if 'cvssMetricV31' in metrics else {}
        cvss_v2 = metrics.get('cvssMetricV2', [{}])[0].get('cvssData', {}) if 'cvssMetricV2' in metrics else {}
        has_exploit = any('Exploit' in ref.get('tags', []) for ref in cve.get('references', []))
        cisa_kev = cve.get('cisaExploitAdd') is not None
        vuln_status = cve.get('vulnStatus', 'UNKNOWN')
        return {
            'cve_id': cve.get('id'),
            'vulnStatus': vuln_status,
            'cvss_v31': {'score': cvss_v31.get('baseScore', 0), 'severity': cvss_v31.get('baseSeverity', 'UNKNOWN'), 'vector': cvss_v31.get('vectorString', '')},
            'cvss_v2': {'score': cvss_v2.get('baseScore', 0), 'vector': cvss_v2.get('vectorString', '')},
            'has_known_exploit': has_exploit,
            'cisa_kev': cisa_kev
        }

class ThreatIntelligenceEnricher:
    def __init__(self, nvd: NVDIntegration):
        self.nvd = nvd
        self.local_mem_cache = {} 
        self.cache_lock = threading.Lock()

    def enrich_vulnerability(self, vuln_data: dict) -> dict:
        cve_ids = vuln_data.get('CVE_IDS', [])
        if isinstance(cve_ids, str): cve_ids = [c.strip() for c in cve_ids.split(',') if c.strip()]
        enriched_intel = {'max_cvss': 0.0, 'has_exploit': False, 'cisa_kev': False, 'cvss_vector': '', 'risk_factors': []}
        
        for cve_id in cve_ids[:3]: 
            cve_data = None
            with self.cache_lock: cve_data = self.local_mem_cache.get(cve_id)
            if not cve_data:
                cve_data = self.nvd.get_cve_details(cve_id)
                if cve_data:
                    with self.cache_lock: self.local_mem_cache[cve_id] = cve_data
            if cve_data:
                score = max(cve_data['cvss_v31']['score'], cve_data['cvss_v2']['score'])
                if score > enriched_intel['max_cvss']:
                    enriched_intel['max_cvss'] = score
                    enriched_intel['cvss_vector'] = cve_data['cvss_v31']['vector'] or cve_data['cvss_v2']['vector']
                if cve_data['has_known_exploit']: enriched_intel['has_exploit'] = True
                if cve_data['cisa_kev']: enriched_intel['cisa_kev'] = True
        return enriched_intel

# --- CLASSE DE INTEGRAÇÃO TENABLE ---
class TenableIntegration:
    def __init__(self, access_key, secret_key, logger_func):
        self.base_url = "https://cloud.tenable.com"
        masked_acc = f"{access_key[:4]}...{access_key[-4:]}" if access_key and len(access_key) > 8 else "INVALID"
        logger_func(f"DEBUG: Configurando API Tenable com chaves (Access: {masked_acc})")
        self.headers = {
            "X-ApiKeys": f"accessKey={access_key};secretKey={secret_key}",
            "User-Agent": "VulnManagerAI/ExportClient",
            "Content-Type": "application/json",
            "Accept": "application/json"
        }
        self.log = logger_func

    def get_full_inventory(self, progress_callback=None):
        self.log(">>> [Tenable Inventory] Iniciando exportação de ativos (Lote Aumentado)...")
        url = f"{self.base_url}/assets/export"
        
        # Aumentado para 5000 para evitar fragmentação excessiva
        payload = {"chunk_size": 5000}
        
        try:
            r = requests.post(url, headers=self.headers, json=payload, timeout=30)
            if r.status_code != 200:
                self.log(f"Erro no Request Asset Export: {r.status_code} - {r.text}")
                return []
            
            export_uuid = r.json().get('export_uuid')
            
            if progress_callback: progress_callback(0.15, "Processando Inventário (Aguarde)...")
            
            status_url = f"{self.base_url}/assets/export/{export_uuid}/status"
            available_chunks = []
            
            while True:
                r = requests.get(status_url, headers=self.headers, timeout=30)
                if r.status_code == 200:
                    status_data = r.json()
                    status = status_data.get('status')
                    if status == 'FINISHED':
                        available_chunks = status_data.get('chunks_available', [])
                        break
                    elif status == 'ERROR':
                        self.log("Tenable retornou erro no processamento do Export.")
                        return []
                time.sleep(3)

            self.log(f"DEBUG: Baixando {len(available_chunks)} chunks de ativos...")
            all_assets = []
            
            for idx, chunk_id in enumerate(available_chunks):
                try:
                    if progress_callback: 
                        progress_callback(0.2, f"Baixando Inventário Chunk {idx+1}/{len(available_chunks)}...")
                    
                    chunk_url = f"{self.base_url}/assets/export/{export_uuid}/chunks/{chunk_id}"
                    r = requests.get(chunk_url, headers=self.headers, timeout=120)
                    
                    if r.status_code == 200:
                        assets_in_chunk = r.json()
                        all_assets.extend(assets_in_chunk)
                    else:
                        self.log(f"   > ERRO ao baixar Chunk {idx+1}: Status {r.status_code}")
                except Exception as e_chunk:
                    self.log(f"   > EXCEPTION no Chunk {idx+1}: {e_chunk}")
            
            return all_assets

        except Exception as e:
            self.log(f"ERRO no Asset Export Geral: {e}")
            return []

    def get_as_csv_string_qualys_format(self, days_back=None, progress_callback=None):
        self.log("="*50)
        
        # PASSO 1: Inventário (Tags)
        inventory_data_raw = self.get_full_inventory(progress_callback)
        
        uuid_to_tags = {}
        for asset in inventory_data_raw:
            aid = asset.get('id')
            tags_raw = asset.get('tags', [])
            tags_formatted = " | ".join([f"{t.get('key')}:{t.get('value')}" for t in tags_raw])
            if aid:
                uuid_to_tags[aid] = tags_formatted

        # PASSO 2: Vulnerabilidades
        self.log(f">>> [Tenable Export] Iniciando Export de Vulnerabilidades...")
        if progress_callback: progress_callback(0.3, "Iniciando Export Vulns...")

        export_uuid = self._request_export(days_back)
        if not export_uuid: return None, {}, [], {}, {}
        
        available_chunks = self._wait_for_export(export_uuid)
        if not available_chunks: return None, {}, [], {}, {}
        
        all_vulns = self._download_chunks(export_uuid, available_chunks, progress_callback)
        self.log(f">>> [Tenable Export] Vulns baixadas. Total: {len(all_vulns)}")
        
        return self._convert_to_qualys_structures(all_vulns, uuid_to_tags)

    def _request_export(self, days_back=None):
        url = f"{self.base_url}/vulns/export"
        if days_back: start_date = int((datetime.now() - timedelta(days=days_back)).timestamp())
        else: start_date = int((datetime.now() - timedelta(days=365*15)).timestamp())

        # Aumentado num_assets para reduzir fragmentação
        payload = { 
            "num_assets": 2500,
            "filters": { 
                "state": ["OPEN", "REOPENED", "FIXED"], 
                "severity": ["low", "medium", "high", "critical"],
                "last_found": start_date,
                "severity_modification_type": ["ACCEPTED"] 
            },
            "include_unlicensed": True 
        }
        try:
            r = requests.post(url, headers=self.headers, json=payload, timeout=30)
            if r.status_code == 200:
                return r.json().get('export_uuid')
            return None
        except Exception as e:
            self.log(f"EXCEPTION no Request Export: {e}"); return None

    def _wait_for_export(self, export_uuid):
        url = f"{self.base_url}/vulns/export/{export_uuid}/status"
        start_time = time.time()
        while True:
            if time.time() - start_time > 600: return []
            try:
                r = requests.get(url, headers=self.headers, timeout=30)
                if r.status_code == 200:
                    status = r.json().get('status')
                    if status == 'FINISHED': return r.json().get('chunks_available', [])
                    elif status == 'ERROR': return []
                time.sleep(5)
            except: time.sleep(5)

    def _download_chunks(self, export_uuid, chunks, progress_callback=None):
        full_data = []
        total_chunks = len(chunks)
        for index, chunk_id in enumerate(chunks):
            url = f"{self.base_url}/vulns/export/{export_uuid}/chunks/{chunk_id}"
            try:
                if progress_callback:
                    pct = 0.4 + (0.3 * ((index + 1) / total_chunks))
                    progress_callback(pct, f"Baixando Vulns Chunk {index + 1}/{total_chunks}")
                r = requests.get(url, headers=self.headers, timeout=120)
                if r.status_code == 200: full_data.extend(r.json())
            except: pass
        return full_data

    def _convert_to_qualys_structures(self, export_data, uuid_to_tags):
        self.log("DEBUG: Gerando CSV Intermediário (Mapeando Portas e Protocolos)...")
        output = StringIO()
        writer = csv.writer(output)
        
        # CABEÇALHO INTERMEDIÁRIO (Adicionando os campos que faltavam)
        writer.writerow([
            "QID", "IP", "DNS", "NetBIOS", "OS", "HOST ID", "STATUS", "Severity", 
            "First Found", "Last Detected",
            "Tenable_Plugin_ID", "VPR_Score", "CVSS_v3_Base", "Risk_Factor", "CVSS_v2_Base",
            "PORT", "PROTOCOL", "FQDN", "SSL", "SERVICE", "INSTANCE"  # <--- NOVOS CAMPOS
        ])
        
        kb_details = {}
        assets_seen = {} 
        sev_map = {'info': 1, 'low': 2, 'medium': 3, 'high': 4, 'critical': 5}
        
        for item in export_data:
            try:
                plugin = item.get('plugin', {})
                asset = item.get('asset', {})
                
                # --- DADOS PADRÃO ---
                plugin_id = str(plugin.get('id', ''))
                plugin_name = plugin.get('name', 'Unknown')
                severity = sev_map.get(item.get('severity', 'info').lower(), 1)
                
                ip_addr = asset.get('ipv4', '0.0.0.0')
                if not ip_addr: ip_addr = asset.get('ipv6', '0.0.0.0')
                
                hostname = asset.get('hostname', '')
                uuid = asset.get('uuid', '')
                
                # Extração de OS
                os_raw = asset.get('operating_system', ['Unknown'])
                os_name = os_raw[0] if isinstance(os_raw, list) and len(os_raw) > 0 else str(os_raw)
                
                first_found = item.get('first_found', datetime.now().isoformat())
                last_found = item.get('last_found', datetime.now().isoformat())
                state = 'Active' if item.get('state', 'OPEN') in ['OPEN', 'REOPENED'] else 'Fixed'
                tags_str = uuid_to_tags.get(uuid, "Tenable Asset")

                # --- NOVAS EXTRAÇÕES (SOLICITADAS) ---
                port = str(item.get('port', '0'))
                if port == '0': port = '' # Limpa se for zero
                
                protocol = str(item.get('protocol', 'TCP')).upper()
                service = str(item.get('service_name', ''))
                
                # Tenable: FQDN geralmente é o hostname ou uma lista em 'fqdn'
                fqdn = asset.get('fqdn', '')
                if isinstance(fqdn, list) and len(fqdn) > 0: fqdn = fqdn[0]
                elif not fqdn: fqdn = hostname # Fallback para hostname
                
                # Tenable não tem campo booleano "SSL", mas podemos inferir ou deixar vazio.
                # Se o serviço tem 'ssl' no nome ou porta 443, podemos sugerir, mas melhor deixar raw.
                ssl_val = "Yes" if 'ssl' in service.lower() or port in ['443', '8443'] else ""
                
                # Instância (Geralmente para DBs, Tenable joga no output, mas vamos tentar pegar de metadados se existir)
                instance = "" # Difícil extrair sem parsing de output do plugin

                # --- VPR & RISCO ---
                vpr_score = ''
                if 'vpr_score' in item: vpr_score = item['vpr_score']
                elif 'vpr' in item and isinstance(item['vpr'], dict): vpr_score = item['vpr'].get('score', '')
                elif 'vpr' in plugin and isinstance(plugin['vpr'], dict): vpr_score = plugin['vpr'].get('score', '')
                elif 'vpr_score' in plugin: vpr_score = plugin['vpr_score']
                vpr_score = str(vpr_score) if vpr_score is not None else ''

                cvss3 = plugin.get('cvss3_base_score', '')
                if not cvss3: cvss3 = item.get('cvss3_base_score', '')
                
                cvss2 = plugin.get('cvss_base_score', '')
                if not cvss2: cvss2 = item.get('cvss_base_score', '')

                risk_factor = plugin.get('risk_factor', '')
                if not risk_factor: risk_factor = item.get('risk_factor', '')

                # --- ESCRITA DA LINHA ---
                writer.writerow([
                    plugin_id, ip_addr, hostname, "", os_name, uuid, state, severity, 
                    first_found, last_found,
                    plugin_id, vpr_score, cvss3, risk_factor, cvss2,
                    port, protocol, fqdn, ssl_val, service, instance # <--- CAMPOS ADICIONADOS
                ])
                
                # Popula KB
                if plugin_id not in kb_details:
                    kb_details[plugin_id] = {
                        'TITLE': plugin_name, 
                        'CATEGORY': plugin.get('family', 'General'),
                        'SOLUTION': plugin.get('solution', 'Consulte o fabricante.'),
                        'CVE_IDS': ",".join(plugin.get('cve', [])) if 'cve' in plugin else ''
                    }

                if uuid and uuid not in assets_seen:
                    assets_seen[uuid] = {
                        'Host ID': uuid, 'IP Address': ip_addr, 'DNS Name': hostname,
                        'Operating System': os_name, 'Tags': tags_str
                    }
            except Exception: 
                continue
        
        tags_by_id = {uid: d['Tags'] for uid, d in assets_seen.items()}
        tags_by_ip = {d['IP Address']: d['Tags'] for d in assets_seen.values() if d['IP Address']}


@dataclass
class SegmentContext:
    name: str
    regulations: list
    critical_assets: list
    business_impact_examples: list
    risk_tolerance: str
    key_concerns: list

class PromptEngineer:
    def __init__(self):
        self.segment_contexts = self._load_segment_contexts()

    def _load_segment_contexts(self) -> Dict[str, SegmentContext]:
        return {
            'financeiro': SegmentContext(name="Setor Financeiro & Bancário", regulations=['PCI DSS 4.0', 'Resolução BCB 4.893', 'LGPD', 'SOX', 'Swift CSP'], critical_assets=['Core Banking', 'Chaves PIX', 'Internet Banking', 'Mobile Banking', 'Gateways de Pagamento', 'Dados de Cartão (PAN)'], business_impact_examples=['Fraude financeira direta e perda de liquidez', 'Vazamento de dados sigilosos e quebra de sigilo bancário', 'Indisponibilidade de transações (PIX/TED) e multas do BACEN', 'Danos reputacionais irreversíveis no mercado'], risk_tolerance='Crítica (Zero Trust)', key_concerns=['Fraude', 'Compliance Regulatório', 'Disponibilidade 24x7', 'Segurança de Transações']),
            'saude': SegmentContext(name="Setor de Saúde & Life Sciences", regulations=['LGPD', 'HIPAA', 'Resoluções CFM/Anvisa'], critical_assets=['Prontuário Eletrônico (PEP)', 'PACS/RIS (Imagens)', 'Equipamentos Médicos (IoMT)', 'Dados Laboratoriais', 'Fórmulas Farmacêuticas'], business_impact_examples=['Vazamento de dados sensíveis de pacientes (PHI)', 'Sequestro de dados (Ransomware) paralisando triagem', 'Risco direto à vida por falha em equipamentos conectados', 'Processos éticos e judiciais'], risk_tolerance='Muito Baixa', key_concerns=['Privacidade do Paciente', 'Segurança Hospitalar', 'Disponibilidade de Sistemas Críticos']),
            'varejo': SegmentContext(name="Setor de Varejo & E-commerce", regulations=['PCI DSS', 'LGPD', 'CDC'], critical_assets=['Plataforma E-commerce', 'PDV (Ponto de Venda)', 'ERP', 'Banco de Dados CRM', 'Gateway de Frete'], business_impact_examples=['Skimming de cartão de crédito no checkout', 'Indisponibilidade do site em datas sazonais (Black Friday)', 'Vazamento de base de clientes para concorrentes', 'Perda imediata de receita por hora parada'], risk_tolerance='Média', key_concerns=['Uptime em alta demanda', 'Proteção de dados de cartão', 'Experiência do Cliente']),
            'industria': SegmentContext(name="Setor Industrial & Manufatura", regulations=['NR-12', 'ISO 27001', 'IEC 62443'], critical_assets=['Sistemas SCADA', 'PLCs', 'Redes OT', 'Sistemas MES', 'Historiadores de Dados', 'Robôs Industriais'], business_impact_examples=['Parada não planejada da linha de produção', 'Alteração de parâmetros de qualidade do produto', 'Danos físicos a equipamentos ou colaboradores', 'Espionagem industrial e roubo de propriedade intelectual'], risk_tolerance='Baixa', key_concerns=['Convergência IT/OT', 'Segurança Física', 'Continuidade Operacional']),
            'energia': SegmentContext(name="Setor de Energia & Utilities", regulations=['Resoluções ANEEL/ANP', 'ONS', 'Lei de Infraestrutura Crítica'], critical_assets=['Centros de Operação (COS)', 'Sistemas de Supervisão', 'Relés de Proteção', 'Medidores Inteligentes', 'Rede de Distribuição'], business_impact_examples=['Blackout ou interrupção de fornecimento', 'Danos à infraestrutura crítica nacional', 'Impacto ambiental por falha de controle', 'Terrorismo cibernético'], risk_tolerance='Zero', key_concerns=['Segurança Nacional', 'Alta Disponibilidade', 'Sistemas Legados']),
            'governo': SegmentContext(name="Setor Público & Governo", regulations=['LGPD', 'LAI', 'Normativas GSI/NC', 'Lei de Segurança Nacional'], critical_assets=['Bases de Dados do Cidadão', 'Portais de Serviços (Gov.br)', 'Sistemas de Arrecadação', 'Infraestrutura de Cidades Inteligentes'], business_impact_examples=['Exposição massiva de dados de cidadãos', 'Interrupção de serviços públicos essenciais', 'Perda de confiança nas instituições', 'Ataques de Hacktivismo'], risk_tolerance='Baixa', key_concerns=['Soberania Digital', 'Privacidade do Cidadão', 'Resiliência']),
            'tecnologia': SegmentContext(name="Setor de Tecnologia & Telecom", regulations=['ISO 27001', 'SOC 2', 'GDPR/LGPD', 'Marco Civil da Internet'], critical_assets=['Código Fonte (IP)', 'Pipelines CI/CD', 'Infraestrutura Cloud', 'Dados de Assinantes', 'Backbones de Rede'], business_impact_examples=['Ataques de Supply Chain afetando clientes', 'Vazamento de credenciais de acesso privilegiado', 'DDoS massivo contra infraestrutura', 'Comprometimento de propriedade intelectual'], risk_tolerance='Média-Baixa', key_concerns=['Segurança de Aplicação', 'Cloud Security', 'Gerenciamento de Identidade']),
            'logistica': SegmentContext(name="Setor de Logística & Transportes", regulations=['LGPD', 'Normas ANTT/ANAC', 'AEO'], critical_assets=['Sistemas de Rastreamento', 'Gestão de Frota', 'WMS (Warehouse Management)', 'Manifestos de Carga'], business_impact_examples=['Interrupção da cadeia de suprimentos', 'Roubo de carga facilitado por dados vazados', 'Atrasos críticos em entregas', 'Perda de visibilidade da frota'], risk_tolerance='Média', key_concerns=['Integridade da Cadeia', 'IoT/Rastreamento', 'Disponibilidade']),
            'educacao': SegmentContext(name="Setor de Educação & Ensino", regulations=['LGPD', 'MEC'], critical_assets=['Plataformas EAD/LMS', 'Dados Acadêmicos e Financeiros de Alunos', 'Propriedade Intelectual de Pesquisas'], business_impact_examples=['Vazamento de dados de menores de idade', 'Interrupção de aulas e provas online', 'Perda de pesquisas científicas não publicadas', 'Danos à reputação institucional'], risk_tolerance='Média', key_concerns=['Privacidade de Alunos', 'Continuidade de Aulas', 'Segurança de Dados de Pesquisa']),
            'juridico': SegmentContext(name="Setor Jurídico & Consultoria", regulations=['LGPD', 'Estatuto da Advocacia', 'ISO 27001'], critical_assets=['Dossiês de Clientes', 'E-mails e Comunicações Sigilosas', 'Sistemas de Gestão de Processos', 'Propriedade Intelectual'], business_impact_examples=['Quebra de sigilo advogado-cliente', 'Espionagem industrial contra clientes', 'Perda de credibilidade e processos disciplinares', 'Ransomware em dados de litígio'], risk_tolerance='Baixa', key_concerns=['Confidencialidade Extrema', 'Controle de Acesso', 'Prevenção de Vazamento (DLP)']),
            'midia': SegmentContext(name="Setor de Mídia & Entretenimento", regulations=['LGPD', 'Direitos Autorais', 'Classificação Indicativa'], critical_assets=['Conteúdo Inédito (Masters)', 'Plataformas de Streaming', 'Dados de Assinantes', 'Canais de Transmissão ao Vivo'], business_impact_examples=['Vazamento de conteúdo antes do lançamento (Spoilers/Pirataria)', 'Interrupção de transmissões ao vivo', 'Sequestro de canais sociais', 'DDoS em estreias'], risk_tolerance='Média', key_concerns=['Proteção de Conteúdo (DRM)', 'Alta Disponibilidade', 'Segurança de Contas'])
        }
        
    def _identify_segment(self, segment_selection: str) -> str:
        mapping = {
            "Financeiro - Bancos & Investimentos": "financeiro", "Financeiro - Meios de Pagamento & Fintechs": "financeiro", "Financeiro - Seguradoras": "financeiro",
            "Varejo - E-commerce (Comércio Eletrônico)": "varejo", "Varejo - Redes de Lojas Físicas & Supermercados": "varejo",
            "Saúde - Hospitais & Clínicas": "saude", "Saúde - Laboratórios de Diagnóstico": "saude", "Saúde - Indústria Farmacêutica": "saude",
            "Tecnologia - Desenvolvimento de Software (SaaS)": "tecnologia", "Tecnologia - Data Centers & Cloud Providers": "tecnologia", "Tecnologia - Telecomunicações & ISP": "tecnologia",
            "Indústria - Manufatura & Fábricas": "industria", "Indústria - Automotiva": "industria", "Indústria - Agronegócio": "industria",
            "Energia - Geração & Distribuição Elétrica": "energia", "Energia - Óleo, Gás & Petroquímica": "energia",
            "Logística - Transporte & Portos": "logistica", "Governo - Federal/Estadual (Dados Cidadão)": "governo",
            "Educação - Universidades & Ensino a Distância": "educacao", "Serviços Jurídicos & Advocacia": "juridico", "Consultoria & Auditoria": "juridico", "Mídia, Entretenimento & Streaming": "midia"
        }
        return mapping.get(segment_selection, 'varejo')

    def _format_candidates(self, candidates: list) -> str:
        if not candidates: return "Nenhuma vulnerabilidade crítica detectada."
        lines = []
        for i, v in enumerate(candidates[:20], 1):
            title = v.get('Vuln_TITLE', 'N/A')
            hosts = v.get('Hosts_Count', 0)
            sev = v.get('Max_Sev', 0)
            risk = v.get('Risk_Sum', 0)
            solution = v.get('Vuln_SOLUTION', 'Consulte o vendor.')[:100] + "..."
            lines.append(f"{i}. **{title}**")
            lines.append(f"   - Hosts Afetados: {hosts} | Severidade: {sev} | Risk Score: {risk:.0f}")
            lines.append(f"   - Resumo Solução: {solution}")
            lines.append("")
        return '\n'.join(lines)

    def generate_analysis_prompt(self, summary: dict, segment: str, period: str, start_date_str: str = "") -> str:
        segment_key = self._identify_segment(segment)
        context = self.segment_contexts.get(segment_key, self.segment_contexts['varejo'])
        status = summary['status_counts']
        active = status['Active'] + status['New'] + status['Re-Opened']
        candidates_str = self._format_candidates(summary.get('top_remediation', []))
        crit_assets_str = "\n".join([f"- {asset}" for asset in context.critical_assets])
        regs_str = ", ".join(context.regulations)
        
        maturity_context = ""
        if start_date_str:
            try:
                start_dt = datetime.strptime(start_date_str, "%d/%m/%Y")
                months_diff = (datetime.now() - start_dt).days // 30
                maturity_context = f"\nIMPORTANTE: A operação de Gestão de Vulnerabilidades iniciou em {start_date_str} (aprox. {months_diff} meses de operação)."
            except: pass

        prompt = f"""
        ATUE COMO UM CISO SÊNIOR ESPECIALIZADO NO: {context.name}. (Porém utilize linguagem simples, não seja tão engessado, não precisa parecer um documento ou e-mail, porém não tão coloquial)
        
        ## 1. CONTEXTO DO NEGÓCIO E REGULATÓRIO
        - **Regulamentações Chave**: {regs_str}
        - **Tolerância a Risco**: {context.risk_tolerance}
        - **Maiores Preocupações**: {', '.join(context.key_concerns)}
        {maturity_context}
        
        ## 2. ATIVOS CRÍTICOS (O QUE DEVEMOS PROTEGER A TODO CUSTO)
        {crit_assets_str}
        
        ## 3. DADOS TÉCNICOS DO RELATÓRIO ({period})
        - **Total Ativas**: {active} (Novas + Reabertas + Ativas)
        - **Total Corrigidas**: {status['Fixed']}
        - **Legado Crítico (>90 dias)**: {summary['smart_insights']['old_crit_count']} ativos vulneráveis.
        
        ## 4. TOP VULNERABILIDADES DETECTADAS (CANDIDATAS)
        {candidates_str}

        ## SUA MISSÃO (RELATÓRIO ESTRATÉGICO)
        
        ### A. ANÁLISE DE RISCO CONTEXTUALIZADA
        Analise os dados acima considerando o **impacto de negócio** específico para {context.name}.
        Cite exemplos reais de como essas falhas poderiam afetar os **Ativos Críticos** listados (Ex: {context.business_impact_examples[0]}).
        
        ### B. PRIORIZAÇÃO DE ELITE (TOP 5)
        Das candidatas listadas, selecione APENAS AS 5 MAIS PERIGOSAS para este setor específico.
        Formato Obrigatório para cada uma das 5:
        1. **[Nome da Vuln]** (QID: Se houver) - (Hosts afetados)
           - **Impacto no Negócio**: [Explique em 1 frase por que isso para o setor {segment_key}]
           - **Conformidade**: [Violação potencial de qual norma? Ex: {context.regulations[0]}]
           - **Ação Imediata**: [Ação técnica resumida]

        ### C. CONCLUSÃO EXECUTIVA
        Escreva um parágrafo final para o Board de analistas.
        """
        return prompt

class ReportChartGenerator:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        try: plt.style.use('seaborn-v0_8-whitegrid')
        except: plt.style.use('ggplot')

    def create_executive_summary_chart(self, summary: dict) -> str:
        fig = plt.figure(figsize=(16, 12))
        gs = GridSpec(3, 3, figure=fig, hspace=0.4, wspace=0.3)

        # 1. Gauge de Risco Geral
        ax1 = fig.add_subplot(gs[0, 0])
        risk_norm = min(summary.get('risk_score_norm', 50), 100)
        self._draw_risk_gauge(ax1, risk_norm)

        # 2. Pizza de Status
        ax2 = fig.add_subplot(gs[0, 1])
        status = summary['status_counts']
        sizes = [status.get('Active', 0), status.get('Fixed', 0), status.get('New', 0), status.get('Re-Opened', 0)]
        labels = ['Ativas', 'Corrigidas', 'Novas', 'Reabertas']
        pie_data = [(s, l) for s, l in zip(sizes, labels) if s > 0]
        if pie_data:
            sizes_f, labels_f = zip(*pie_data)
            colors = ['#FF6B6B', '#4ECDC4', '#FFE66D', '#95E1D3']
            ax2.pie(sizes_f, labels=labels_f, colors=colors[:len(sizes_f)], autopct='%1.1f%%', startangle=90, explode=[0.05]*len(sizes_f))
        ax2.set_title('Distribuição por Status', fontweight='bold')

        # 3. Barras de Severidade
        ax3 = fig.add_subplot(gs[0, 2])
        severities = summary.get('severity_counts', {})
        sev_labels = ['Crítico (5)', 'Alto (4)', 'Médio (3)', 'Baixo (2)', 'Info (1)']
        sev_values = [severities.get(5, 0), severities.get(4, 0), severities.get(3, 0), severities.get(2, 0), severities.get(1, 0)]
        sev_colors = ['#8B0000', '#FF4500', '#FFA500', '#FFD700', '#90EE90']
        bars = ax3.barh(sev_labels[::-1], sev_values[::-1], color=sev_colors[::-1])
        ax3.bar_label(bars)
        ax3.set_title('Vulns por Severidade', fontweight='bold')

        # 4. Histograma de Envelhecimento (Aging)
        ax4 = fig.add_subplot(gs[1, :2])
        aging = summary.get('aging_buckets', {})
        buckets = ['0-30 Dias', '31-60 Dias', '61-90 Dias', '90+ Dias (Legado)']
        vals = [aging.get('0-30', 0), aging.get('31-60', 0), aging.get('61-90', 0), aging.get('90+', 0)]
        colors_age = ['#4ECDC4', '#FFE66D', '#FF6B6B', '#2C3E50']
        ax4.bar(buckets, vals, color=colors_age, alpha=0.8)
        ax4.set_title('Envelhecimento das Vulnerabilidades (Aging)', fontweight='bold')
        ax4.grid(axis='y', linestyle='--', alpha=0.7)

        # 5. MTTR por Severidade
        ax5 = fig.add_subplot(gs[1, 2])
        mttr = summary.get('avg_time_to_fix', {})
        if mttr:
            mttr_labels = [f'S{k}' for k in sorted(mttr.keys(), reverse=True)]
            mttr_values = [mttr[k] for k in sorted(mttr.keys(), reverse=True)]
            ax5.bar(mttr_labels, mttr_values, color='steelblue')
            ax5.axhline(y=30, color='red', linestyle='--', linewidth=1, label='SLA 30d')
            ax5.legend(fontsize=8)
        else:
            ax5.text(0.5, 0.5, "Sem dados de correção", ha='center')
        ax5.set_title('MTTR Médio (Dias)', fontweight='bold')

        # 6. Top 10 Vulnerabilidades
        ax6 = fig.add_subplot(gs[2, :])
        top_vulns = summary.get('top_remediation', [])[:10]
        if top_vulns:
            titles = [v['Vuln_TITLE'][:60] + '...' for v in top_vulns]
            risks = [v['Risk_Sum'] for v in top_vulns]
            y_pos = np.arange(len(titles))
            ax6.barh(y_pos, risks, align='center', color='#C0392B')
            ax6.set_yticks(y_pos)
            ax6.set_yticklabels(titles, fontsize=9)
            ax6.invert_yaxis()
            ax6.set_xlabel('Risk Score Acumulado')
            ax6.set_title('Top 10 Vulnerabilidades Críticas (Prioridade de Remediação)', fontweight='bold')
        else:
            ax6.text(0.5, 0.5, "Nenhuma vulnerabilidade crítica encontrada", ha='center')

        plt.suptitle(f'Panorama de Segurança Cibernética\nGerado em: {datetime.now().strftime("%d/%m/%Y")}', fontsize=16, fontweight='bold', y=0.95)

        path = os.path.join(self.output_dir, 'executive_dashboard_v2.png')
        plt.savefig(path, dpi=150, bbox_inches='tight')
        plt.close()
        return path

    def _draw_risk_gauge(self, ax, score):
        ax.set_xlim(-1.5, 1.5); ax.set_ylim(-0.5, 1.5); ax.axis('off')
        theta = np.linspace(np.pi, 0, 100)
        r = 1; x = r * np.cos(theta); y = r * np.sin(theta)
        colors = [('green', 0, 33), ('gold', 33, 66), ('red', 66, 100)]
        for color, start_p, end_p in colors:
            start_idx = int(start_p * len(theta) / 100)
            end_idx = int(end_p * len(theta) / 100)
            if end_idx >= len(theta): end_idx = len(theta) -1
            ax.plot(x[start_idx:end_idx], y[start_idx:end_idx], color=color, linewidth=25, solid_capstyle='butt')
        angle = np.pi - (score / 100) * np.pi
        ax.arrow(0, 0, 0.8 * np.cos(angle), 0.8 * np.sin(angle), head_width=0.1, head_length=0.1, fc='black', ec='black', width=0.02)
        circle = mpatches.Circle((0, 0), 0.1, color='black'); ax.add_patch(circle)
        ax.text(0, -0.4, f'{score:.0f}/100', ha='center', va='center', fontsize=20, fontweight='bold', color='#333')
        status_text = "CRÍTICO" if score > 75 else "ALTO" if score > 50 else "MÉDIO" if score > 25 else "BAIXO"
        color_text = "red" if score > 75 else "orange" if score > 50 else "gold" if score > 25 else "green"
        ax.text(0, 0.5, status_text, ha='center', va='center', fontsize=10, fontweight='bold', color='white', bbox=dict(facecolor=color_text, edgecolor='none', alpha=0.7, boxstyle='round'))
        ax.set_title('Índice de Risco do Ambiente', fontsize=11, y=-0.1)

# --- FLASK APP SETUP (RICH DASHBOARD) ---
flask_app = Flask(__name__)

DASHBOARD_HTML = """
<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>VulnManager AI | Dashboard Executivo</title>
    <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <style>
        :root { --bg-main: #0f172a; --bg-card: #1e293b; --text-primary: #f8fafc; --text-secondary: #94a3b8; --accent: #38bdf8; --danger: #ef4444; --warning: #f59e0b; --success: #10b981; }
        body { background-color: var(--bg-main); color: var(--text-primary); font-family: 'Segoe UI', sans-serif; overflow-x: hidden; }
        .navbar { background-color: #1e293b; border-bottom: 1px solid #334155; }
        .card { background-color: var(--bg-card); border: 1px solid #334155; border-radius: 12px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.2); margin-bottom: 20px; }
        .metric-box { padding: 1.5rem; position: relative; overflow: hidden; }
        .metric-value { font-size: 2rem; font-weight: 700; margin-bottom: 2px; }
        .metric-label { font-size: 0.8rem; text-transform: uppercase; color: var(--text-secondary); letter-spacing: 0.5px; font-weight: 600; }
        .icon-bg { position: absolute; top: 15px; right: 15px; font-size: 3rem; opacity: 0.08; transform: rotate(-10deg); }
        .table-dark { background-color: var(--bg-card); color: var(--text-secondary); --bs-table-bg: transparent; }
        .badge-sev-5 { background-color: #7f1d1d; color: #fca5a5; border: 1px solid #ef4444; }
        .badge-sev-4 { background-color: #431407; color: #fdba74; border: 1px solid #f97316; }
        .badge-sev-3 { background-color: #422006; color: #fde047; border: 1px solid #eab308; }
        
        /* Scrollbar */
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-track { background: var(--bg-main); }
        ::-webkit-scrollbar-thumb { background: #475569; border-radius: 4px; }
        ::-webkit-scrollbar-thumb:hover { background: #64748b; }
    </style>
</head>
<body>
    <nav class="navbar navbar-dark mb-4 px-3 py-3">
        <div class="container-fluid">
            <span class="navbar-brand mb-0 h1 fw-bold"><i class="fas fa-shield-halved me-2 text-info"></i>VulnManager AI <span class="text-secondary fw-normal">| {{ client }}</span></span>
            <span class="text-secondary small"><i class="far fa-clock me-1"></i> Atualizado: {{ updated_at }}</span>
        </div>
    </nav>

    <div class="container-fluid px-4">
        <div class="row mb-2">
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box border-start border-4 border-danger">
                    <div class="metric-value text-white" id="kpi-active">--</div>
                    <div class="metric-label text-danger">Total Ativas</div>
                    <i class="fas fa-bug icon-bg text-danger"></i>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box border-start border-4 border-warning">
                    <div class="metric-value text-white" id="kpi-risk">--</div>
                    <div class="metric-label text-warning">Risco Global</div>
                    <i class="fas fa-radiation icon-bg text-warning"></i>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box border-start border-4 border-info">
                    <div class="metric-value text-white" id="kpi-hosts">--</div>
                    <div class="metric-label text-info">Assets Afetados</div>
                    <i class="fas fa-server icon-bg text-info"></i>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box border-start border-4 border-success">
                    <div class="metric-value text-white" id="kpi-fixed">--</div>
                    <div class="metric-label text-success">Corrigidas (Total)</div>
                    <i class="fas fa-check-double icon-bg text-success"></i>
                </div>
            </div>
        </div>

        <div class="row mb-4">
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box py-3 bg-opacity-10" style="background-color: rgba(239, 68, 68, 0.05);">
                    <div class="d-flex justify-content-between align-items-center">
                        <div>
                            <div class="h4 mb-0 fw-bold text-danger" id="kpi-crit">--</div>
                            <div class="small text-secondary">Críticas (Sev 5)</div>
                        </div>
                        <i class="fas fa-exclamation-triangle text-danger opacity-50 fa-lg"></i>
                    </div>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box py-3 bg-opacity-10" style="background-color: rgba(249, 115, 22, 0.05);">
                    <div class="d-flex justify-content-between align-items-center">
                        <div>
                            <div class="h4 mb-0 fw-bold text-warning" id="kpi-exploit">--</div>
                            <div class="small text-secondary">Exploit Disponível</div>
                        </div>
                        <i class="fas fa-skull text-warning opacity-50 fa-lg"></i>
                    </div>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box py-3">
                    <div class="d-flex justify-content-between align-items-center">
                        <div>
                            <div class="h4 mb-0 fw-bold text-light" id="kpi-age">--</div>
                            <div class="small text-secondary">Média Idade (Dias)</div>
                        </div>
                        <i class="far fa-calendar-alt text-secondary opacity-50 fa-lg"></i>
                    </div>
                </div>
            </div>
            <div class="col-xl-3 col-md-6">
                <div class="card metric-box py-3">
                    <div class="d-flex justify-content-between align-items-center">
                        <div>
                            <div class="h4 mb-0 fw-bold text-light" id="kpi-legacy">--</div>
                            <div class="small text-secondary">Legado (>90 dias)</div>
                        </div>
                        <i class="fas fa-history text-secondary opacity-50 fa-lg"></i>
                    </div>
                </div>
            </div>
        </div>

        <div class="row">
            <div class="col-md-4">
                <div class="card h-100">
                    <div class="card-header bg-transparent border-0 fw-bold text-light d-flex justify-content-between">
                        <span><i class="fas fa-chart-pie me-2 text-info"></i>Severidade</span>
                    </div>
                    <div class="card-body p-1"><div id="chart-severity" style="height: 280px;"></div></div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card h-100">
                    <div class="card-header bg-transparent border-0 fw-bold text-light d-flex justify-content-between">
                        <span><i class="fas fa-chart-bar me-2 text-warning"></i>Aging (Envelhecimento)</span>
                    </div>
                    <div class="card-body p-1"><div id="chart-aging" style="height: 280px;"></div></div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card h-100">
                    <div class="card-header bg-transparent border-0 fw-bold text-light d-flex justify-content-between">
                        <span><i class="fas fa-list-ol me-2 text-danger"></i>Top 10 Recorrência</span>
                    </div>
                    <div class="card-body p-0">
                        <div id="table-top10" class="table-responsive" style="max-height: 280px;">
                            <div class="text-center text-secondary mt-5"><div class="spinner-border" role="status"></div><br>Carregando dados...</div>
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script>
        // Função segura para números (Evita null/undefined na tela)
        const safeNum = (val) => (val === undefined || val === null) ? 0 : val;
        
        async function loadData() {
            try {
                // 1. Fetch Metrics
                const m = await fetch('/api/metrics').then(r => r.json());
                
                // KPIs Linha 1
                document.getElementById('kpi-active').innerText = safeNum(m.active).toLocaleString();
                document.getElementById('kpi-risk').innerText = safeNum(m.risk).toLocaleString();
                document.getElementById('kpi-hosts').innerText = safeNum(m.hosts).toLocaleString();
                document.getElementById('kpi-fixed').innerText = safeNum(m.fixed).toLocaleString();

                // KPIs Linha 2 (Novos - precisam existir no backend, senão mostram 0)
                document.getElementById('kpi-crit').innerText = safeNum(m.critical_count).toLocaleString();
                document.getElementById('kpi-exploit').innerText = safeNum(m.exploit_count).toLocaleString();
                document.getElementById('kpi-age').innerText = safeNum(m.avg_age).toFixed(0);
                document.getElementById('kpi-legacy').innerText = safeNum(m.legacy_count).toLocaleString();

                // 2. Fetch Charts
                const layoutConfig = {
                    paper_bgcolor: 'rgba(0,0,0,0)',
                    plot_bgcolor: 'rgba(0,0,0,0)',
                    font: { color: '#94a3b8' },
                    margin: { t: 20, b: 20, l: 30, r: 30 },
                    showlegend: true,
                    legend: { orientation: 'h', y: -0.1 }
                };

                const config = { responsive: true, displayModeBar: false };

                // Severidade
                const sevData = await fetch('/api/charts/severity').then(r => r.json());
                if(sevData.data) {
                    sevData.layout = {...sevData.layout, ...layoutConfig};
                    Plotly.newPlot('chart-severity', sevData.data, sevData.layout, config);
                }

                // Aging
                const ageData = await fetch('/api/charts/aging').then(r => r.json());
                if(ageData.data) {
                    ageData.layout = {...ageData.layout, ...layoutConfig};
                    Plotly.newPlot('chart-aging', ageData.data, ageData.layout, config);
                }

                // 3. Top 10 Table
                const topData = await fetch('/api/charts/top_freq').then(r => r.json());
                let html = '<table class="table table-dark table-hover table-sm mb-0 align-middle"><thead><tr><th class="ps-3">Vulnerabilidade</th><th class="text-end pe-3">Hosts</th></tr></thead><tbody>';
                
                if (topData.rows && topData.rows.length > 0) {
                    topData.rows.forEach(r => {
                        // Simula uma severidade baseada no nome para cor (apenas visual se não vier do backend)
                        let badgeClass = 'text-secondary';
                        let title = r[0] || 'Desconhecido';
                        let count = r[1] || 0;
                        
                        html += `<tr>
                            <td class="ps-3">
                                <div class="text-truncate fw-500 text-light" style="max-width: 250px;" title="${title}">${title}</div>
                            </td>
                            <td class="text-end pe-3">
                                <span class="badge bg-secondary bg-opacity-25 text-info border border-info border-opacity-25 rounded-pill px-3">${count}</span>
                            </td>
                        </tr>`;
                    });
                } else {
                    html += '<tr><td colspan="2" class="text-center py-4 text-muted">Nenhum dado disponível</td></tr>';
                }
                html += '</tbody></table>';
                document.getElementById('table-top10').innerHTML = html;

            } catch (error) {
                console.error("Erro ao carregar dashboard:", error);
            }
        }

        document.addEventListener('DOMContentLoaded', loadData);
        setInterval(loadData, 60000); // Refresh a cada 60s
    </script>
</body>
</html>
"""

class DashboardGenerator:
    def __init__(self, df: pd.DataFrame):
        self.df = df.copy() if not df.empty else pd.DataFrame()
        
        # 1. Normalização de Colunas (Upper Case para facilitar busca)
        self.col_map = {c.upper().strip(): c for c in self.df.columns}

        # 2. Localização Inteligente de Colunas (Compatível com o novo CSV gerado)
        self.col_sev = self._find_col(['SEVERITY', 'SEV', 'MAX_SEV'])
        self.col_status = self._find_col(['STATUS', 'STATE', 'VULN STATUS'])
        self.col_title = self._find_col(['TITLE', 'VULN_TITLE', 'VULN TITLE', 'QID']) # Fallback para QID se Title faltar
        self.col_ip = self._find_col(['IP ADDRESS', 'IP', 'IPADDRESS'])
        self.col_exploit = self._find_col(['THREAT_EXPLOIT', 'EXPLOIT'])
        self.col_first_found = self._find_col(['FIRST FOUND', 'FIRST_FOUND'])

        # 3. Pré-processamento de Dados
        if not self.df.empty:
            # Severidade (Força Numérico)
            if self.col_sev:
                self.df[self.col_sev] = pd.to_numeric(self.df[self.col_sev], errors='coerce').fillna(1).astype(int)
                self.df['Tech_Risk'] = 4 ** self.df[self.col_sev]
            else:
                self.df['Severity'] = 1
                self.df['Tech_Risk'] = 1
                self.col_sev = 'Severity'

            # Datas e Idade (Aging) - CORREÇÃO CRÍTICA AQUI
            self.df['Age_Days'] = 0 # Valor padrão
            if self.col_first_found:
                # Tenta converter com dayfirst=True (padrão BR/Europeu) e UTC
                self.df[self.col_first_found] = pd.to_datetime(self.df[self.col_first_found], errors='coerce', dayfirst=True, utc=True)
                
                now = pd.Timestamp.now(timezone.utc)
                # Calcula a diferença apenas para datas válidas
                mask_valid = self.df[self.col_first_found].notna()
                self.df.loc[mask_valid, 'Age_Days'] = (now - self.df.loc[mask_valid, self.col_first_found]).dt.days
                
                # Garante que não tenha números negativos
                self.df['Age_Days'] = self.df['Age_Days'].clip(lower=0).fillna(0).astype(int)

    def _find_col(self, candidates):
        """Busca case-insensitive no DataFrame"""
        if self.df.empty: return None
        for cand in candidates:
            cand_upper = cand.upper().strip()
            if cand_upper in self.col_map:
                return self.col_map[cand_upper]
        return None

    # ... (Mantenha os outros métodos get_metrics, create_severity_distribution_chart, etc. iguais)
    
    def get_metrics(self):
       if self.df.empty: 
           return {"active": 0, "risk": 0, "hosts": 0, "fixed": 0, "critical_count": 0, "exploit_count": 0, "avg_age": 0, "legacy_count": 0}
       
       # Filtro de Ativas (Ignora Fixed)
       df_active = self.df
       if self.col_status:
           df_active = self.df[~self.df[self.col_status].astype(str).str.upper().str.contains('FIXED', na=False)]
       
       active_count = len(df_active)
       fixed_count = len(self.df) - active_count
       risk = int(df_active['Tech_Risk'].mean()) if not df_active.empty else 0
       
       hosts = 0
       if self.col_ip: hosts = df_active[self.col_ip].nunique()
           
       # KPIs Linha 2
       crit_count = len(df_active[df_active[self.col_sev] >= 4]) if not df_active.empty and self.col_sev else 0
           
       exploit_count = 0
       if self.col_exploit:
            exploit_count = len(df_active[df_active[self.col_exploit].astype(str).str.upper().isin(['SIM', 'YES', 'TRUE', '1'])])

       avg_age = df_active['Age_Days'].mean() if not df_active.empty else 0
       
       # Legado > 90 dias
       legacy_count = len(df_active[df_active['Age_Days'] > 90]) if not df_active.empty else 0
           
       return {
           "active": active_count, 
           "risk": risk, 
           "hosts": hosts, 
           "fixed": fixed_count,
           "critical_count": crit_count,
           "exploit_count": exploit_count,
           "avg_age": int(avg_age),
           "legacy_count": legacy_count
       }

# --- ROTAS FLASK ---
@flask_app.route('/')
def dashboard_index():
    return render_template_string(DASHBOARD_HTML, client=shared_data.client_name, updated_at=shared_data.updated_at.strftime("%d/%m/%Y %H:%M"))

@flask_app.route('/api/metrics')
def api_metrics():
    dash = DashboardGenerator(shared_data.dataframe)
    return jsonify(dash.get_metrics())

@flask_app.route('/api/charts/<chart_type>')
def get_chart(chart_type):
    dash = DashboardGenerator(shared_data.dataframe)
    if chart_type == 'severity': return dash.create_severity_distribution_chart()
    if chart_type == 'top_freq': return dash.get_top_frequent_data()
    if chart_type == 'aging': return dash.create_aging_chart()
    return jsonify({}) 

def run_flask():
    flask_app.run(port=5000, debug=False, use_reloader=False)

# Iniciando Flask
threading.Thread(target=run_flask, daemon=True).start()

# --- 3. APLICAÇÃO PRINCIPAL (UI + LÓGICA COMPLETA) ---

class ModernQualysApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("VulnManager AI")
        self.geometry("1280x900")
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        self.lista_segmentos = [
            "Financeiro - Bancos & Investimentos", "Financeiro - Meios de Pagamento & Fintechs", "Financeiro - Seguradoras",
            "Varejo - E-commerce (Comércio Eletrônico)", "Varejo - Redes de Lojas Físicas & Supermercados",
            "Saúde - Hospitais & Clínicas", "Saúde - Laboratórios de Diagnóstico", "Saúde - Indústria Farmacêutica",
            "Tecnologia - Desenvolvimento de Software (SaaS)", "Tecnologia - Data Centers & Cloud Providers", "Tecnologia - Telecomunicações & ISP",
            "Indústria - Manufatura & Fábricas", "Indústria - Automotiva", "Indústria - Agronegócio",
            "Energia - Geração & Distribuição Elétrica", "Energia - Óleo, Gás & Petroquímica", "Logística - Transporte & Portos",
            "Governo - Federal/Estadual (Dados Cidadão)", "Educação - Universidades & Ensino a Distância",
            "Serviços Jurídicos & Advocacia", "Consultoria & Auditoria", "Mídia, Entretenimento & Streaming"
        ]
        self.lista_periodos = ["Todo período", "90 dias", "45 dias", "30 dias", "15 dias", "7 dias"]
        self.clientes_encontrados = self.buscar_clientes_env()

        self.sidebar = ctk.CTkFrame(self, width=220, corner_radius=0)
        self.sidebar.grid(row=0, column=0, rowspan=4, sticky="nsew")
        self.sidebar.grid_rowconfigure(10, weight=1)

        self.logo_label = ctk.CTkLabel(self.sidebar, text="🛡️ VulnManager AI", font=ctk.CTkFont(size=22, weight="bold"))
        self.logo_label.grid(row=0, column=0, padx=20, pady=(30, 20))

        self.btn_dashboard = ctk.CTkButton(self.sidebar, text="📊 Dashboard", command=self.show_dashboard, fg_color="transparent", border_width=1, text_color=("gray10", "#DCE4EE"), anchor="w")
        self.btn_dashboard.grid(row=1, column=0, padx=20, pady=5, sticky="ew")

        self.btn_scan = ctk.CTkButton(self.sidebar, text="🔍 Novo Scan & Report", command=self.show_scan_config, fg_color="transparent", border_width=1, text_color=("gray10", "#DCE4EE"), anchor="w")
        self.btn_scan.grid(row=2, column=0, padx=20, pady=5, sticky="ew")
        
        self.btn_web = ctk.CTkButton(self.sidebar, text="🌐 Abrir Dashboard Web", command=self.open_web_dashboard, fg_color="#1F6AA5", text_color="white", hover_color="#144870", anchor="w")
        self.btn_web.grid(row=3, column=0, padx=20, pady=5, sticky="ew")

        self.appearance_mode = ctk.CTkOptionMenu(self.sidebar, values=["Dark", "Light", "System"], command=self.change_appearance)
        self.appearance_mode.grid(row=11, column=0, padx=20, pady=20)

        self.main_frame = ctk.CTkFrame(self, corner_radius=10, fg_color="transparent")
        self.main_frame.grid(row=0, column=1, padx=20, pady=20, sticky="nsew")
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_rowconfigure(1, weight=1)

        self.show_dashboard()

    def clear_main_frame(self):
        for widget in self.main_frame.winfo_children(): widget.destroy()

    def change_appearance(self, mode):
        ctk.set_appearance_mode(mode.lower())

    def open_web_dashboard(self):
        webbrowser.open("http://127.0.0.1:5000")

    def show_dashboard(self):
        self.clear_main_frame()
        title = ctk.CTkLabel(self.main_frame, text="Dashboard Executivo", font=ctk.CTkFont(size=28, weight="bold"))
        title.grid(row=0, column=0, padx=10, pady=(10, 20), sticky="w")
        kpi_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        kpi_frame.grid(row=1, column=0, sticky="nsew")
        kpis = [("Status API IA", "Online" if AI_AVAILABLE else "OFFLINE", "#10b981" if AI_AVAILABLE else "#ef4444"), ("Status Key NVD", "Online" if NVD_API_KEYS else "OFFLINE", "#10b981" if NVD_API_KEYS else "#ef4444"), ("Clientes Ativos", str(len(self.clientes_encontrados)), "#0ea5e9"), ("Versão Engine", "4.0.0", "#6366f1")]
        for i, (label, value, color) in enumerate(kpis):
            card = ctk.CTkFrame(kpi_frame, fg_color=color, corner_radius=12)
            card.grid(row=0, column=i, padx=8, pady=10, sticky="nsew")
            kpi_frame.grid_columnconfigure(i, weight=1)
            ctk.CTkLabel(card, text=value, font=ctk.CTkFont(size=28, weight="bold"), text_color="white").pack(pady=(25, 2))
            ctk.CTkLabel(card, text=label, font=ctk.CTkFont(size=12, weight="normal"), text_color="white").pack(pady=(0, 25))
        info = ctk.CTkLabel(self.main_frame, text="Selecione 'Novo Scan' no menu lateral para iniciar uma nova análise completa.\nApós o scan, clique em 'Abrir Dashboard Web' para visualizar os dados interativos.", font=ctk.CTkFont(size=14))
        info.grid(row=2, column=0, pady=40)

    # FIX 2: Adicionada Barra de Progresso na UI
    def show_scan_config(self):
        self.clear_main_frame()
        title = ctk.CTkLabel(self.main_frame, text="Configuração de Scan & Relatório", font=ctk.CTkFont(size=24, weight="bold"))
        title.grid(row=0, column=0, padx=10, pady=10, sticky="w")
        config_frame = ctk.CTkFrame(self.main_frame)
        config_frame.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        config_frame.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(config_frame, text="Cliente (Env):", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, padx=15, pady=15, sticky="w")
        self.client_combo = ctk.CTkComboBox(config_frame, values=self.clientes_encontrados, width=300)
        self.client_combo.grid(row=0, column=1, padx=15, pady=15, sticky="w")
        if self.clientes_encontrados: self.client_combo.set(self.clientes_encontrados[0])
        ctk.CTkLabel(config_frame, text="Segmento de Negócio:", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, padx=15, pady=15, sticky="w")
        self.segment_combo = ctk.CTkComboBox(config_frame, values=self.lista_segmentos, width=300)
        self.segment_combo.grid(row=1, column=1, padx=15, pady=15, sticky="w")
        self.segment_combo.set(self.lista_segmentos[0])
        ctk.CTkLabel(config_frame, text="Filtro de Período:", font=ctk.CTkFont(weight="bold")).grid(row=2, column=0, padx=15, pady=15, sticky="w")
        self.period_combo = ctk.CTkComboBox(config_frame, values=self.lista_periodos, width=300)
        self.period_combo.grid(row=2, column=1, padx=15, pady=15, sticky="w")
        self.period_combo.set(self.lista_periodos[0])
        ctk.CTkLabel(config_frame, text="Início da Gestão (Opcional):", font=ctk.CTkFont(weight="bold")).grid(row=3, column=0, padx=15, pady=15, sticky="w")
        self.start_date_entry = ctk.CTkEntry(config_frame, width=300, placeholder_text="DD/MM/AAAA (Para cálculo de maturidade)")
        self.start_date_entry.grid(row=3, column=1, padx=15, pady=15, sticky="w")
        ctk.CTkLabel(config_frame, text="Status das Vulns:", font=ctk.CTkFont(weight="bold")).grid(row=4, column=0, padx=15, pady=15, sticky="w")
        chk_frame = ctk.CTkFrame(config_frame, fg_color="transparent")
        chk_frame.grid(row=4, column=1, padx=15, pady=15, sticky="w")
        self.chk_new = ctk.CTkCheckBox(chk_frame, text="New"); self.chk_new.pack(side="left", padx=5); self.chk_new.select()
        self.chk_active = ctk.CTkCheckBox(chk_frame, text="Active"); self.chk_active.pack(side="left", padx=5); self.chk_active.select()
        self.chk_reopened = ctk.CTkCheckBox(chk_frame, text="Re-Opened"); self.chk_reopened.pack(side="left", padx=5); self.chk_reopened.select()
        self.chk_fixed = ctk.CTkCheckBox(chk_frame, text="Fixed"); self.chk_fixed.pack(side="left", padx=5); self.chk_fixed.select()
        ctk.CTkLabel(config_frame, text="Inteligência Artificial:", font=ctk.CTkFont(weight="bold")).grid(row=5, column=0, padx=15, pady=15, sticky="w")
        self.ai_switch = ctk.CTkSwitch(config_frame, text="Gerar Análise (Gemini/Groq)")
        self.ai_switch.grid(row=5, column=1, padx=15, pady=15, sticky="w")
        if AI_AVAILABLE: self.ai_switch.select()
        ctk.CTkLabel(config_frame, text="Threat Intelligence:", font=ctk.CTkFont(weight="bold")).grid(row=6, column=0, padx=15, pady=15, sticky="w")
        self.nvd_switch = ctk.CTkSwitch(config_frame, text="Ativar Enriquecimento NVD (CVE/Exploit)")
        self.nvd_switch.grid(row=6, column=1, padx=15, pady=15, sticky="w")
        if NVD_API_KEYS: self.nvd_switch.select()
        
        # Barra de Progresso
        self.progress_label = ctk.CTkLabel(config_frame, text="Status: Aguardando Início...", text_color="gray")
        self.progress_label.grid(row=7, column=0, columnspan=2, padx=30, pady=(20, 5), sticky="ew")
        
        self.progress_bar = ctk.CTkProgressBar(config_frame, orientation="horizontal", height=15)
        self.progress_bar.grid(row=8, column=0, columnspan=2, padx=30, pady=(0, 20), sticky="ew")
        self.progress_bar.set(0)

        self.btn_execute = ctk.CTkButton(config_frame, text="🚀 INICIAR GERAÇÃO DE RELATÓRIO", font=ctk.CTkFont(size=15, weight="bold"), height=50, corner_radius=10, fg_color="#1F6AA5", hover_color="#144870", command=self.iniciar_thread)
        self.btn_execute.grid(row=9, column=0, columnspan=2, padx=30, pady=(5, 30), sticky="ew")
        log_label = ctk.CTkLabel(self.main_frame, text="Console de Execução:", font=ctk.CTkFont(weight="bold"))
        log_label.grid(row=2, column=0, padx=10, pady=(20,0), sticky="w")
        self.log_box = ctk.CTkTextbox(self.main_frame, height=200, font=("Consolas", 12), border_width=1, border_color="#334155")
        self.log_box.grid(row=3, column=0, padx=10, pady=5, sticky="nsew")
        self.log_box.configure(state="disabled")

    def update_progress(self, val, text):
        try:
            self.after(0, lambda: self.progress_bar.set(val))
            self.after(0, lambda: self.progress_label.configure(text=f"Status: {text}"))
        except: pass

    def buscar_clientes_env(self):
        clientes = set()
        for key in os.environ.keys():
            if key.startswith("QUALYS_USER_"):
                clientes.add(key.replace("QUALYS_USER_", ""))
            if key.startswith("TENABLE_ACCESS_KEY_"):
                clientes.add(key.replace("TENABLE_ACCESS_KEY_", ""))
        return sorted(list(clientes))

    def log(self, msg):
        try:
            timestamp = datetime.now().strftime("[%H:%M:%S]")
            full_msg = f"{timestamp} {msg}\n"
            print(full_msg.strip())
            def _update():
                if hasattr(self, 'log_box'):
                    self.log_box.configure(state="normal")
                    self.log_box.insert("end", full_msg)
                    self.log_box.see("end")
                    self.log_box.configure(state="disabled")
            self.after(0, _update)
        except: pass

    def iniciar_thread(self):
        cliente = self.client_combo.get()
        if not cliente: messagebox.showerror("Erro", "Selecione um cliente."); return
        self.btn_execute.configure(state="disabled", text="PROCESSANDO... AGUARDE")
        self.progress_bar.set(0)
        self.progress_label.configure(text="Iniciando...")
        if hasattr(self, 'log_box'): self.log_box.configure(state="normal"); self.log_box.delete("1.0", "end"); self.log_box.configure(state="disabled")
        status_str = ",".join([s for s in ["New", "Active", "Re-Opened", "Fixed"] if getattr(self, f"chk_{s.lower().replace('-','')}").get()]) or "New,Active,Re-Opened,Fixed"
        threading.Thread(target=self.executar_logica, args=(cliente, status_str, self.segment_combo.get(), self.period_combo.get(), bool(self.ai_switch.get()), self.start_date_entry.get(), bool(self.nvd_switch.get()))).start()

    def finalizar_execucao(self, sucesso=True, msg=""):
        self.btn_execute.configure(state="normal", text="INICIAR GERAÇÃO DE RELATÓRIO")
        self.progress_bar.set(1.0 if sucesso else 0)
        self.progress_label.configure(text=msg)
        if sucesso: self.log(f"SUCESSO: {msg}"); messagebox.showinfo("Concluído", msg)
        else: self.log(f"FALHA: {msg}"); messagebox.showerror("Erro", "Falha no processo. Verifique o log.")

    def perform_ai_analysis(self, summary, segment, period, start_date_str):
        try:
            engineer = PromptEngineer()
            prompt = engineer.generate_analysis_prompt(summary, segment, period, start_date_str)
            if GEMINI_KEY:
                self.log(">>> Conectando ao Google Gemini...")
                genai.configure(api_key=GEMINI_KEY)
                model = genai.GenerativeModel('gemini-2.5-flash')
                response = model.generate_content(prompt)
                return response.text, "Google Gemini"
            if GROQ_KEY:
                self.log(">>> Conectando à Groq Cloud (Llama-3)...")
                client = Groq(api_key=GROQ_KEY)
                chat = client.chat.completions.create(messages=[{"role": "user", "content": prompt}], model="llama-3.3-70b-versatile", temperature=0.3)
                return chat.choices[0].message.content, "Groq (Llama 3)"
            return "Nenhuma chave de API configurada.", "Erro"
        except Exception as e: return f"Erro IA: {str(e)}", "Erro"

    def executing_tenable_logic(self, client_name, nvd_enabled, periodo_str):
        acc = os.getenv(f"TENABLE_ACCESS_KEY_{client_name}")
        sec = os.getenv(f"TENABLE_SECRET_KEY_{client_name}")
        tenable = TenableIntegration(acc, sec, self.log)
        
        days_back = None
        if periodo_str and "Todo" not in periodo_str:
            try:
                days_back = int(periodo_str.split()[0])
            except:
                days_back = 30
        
        # Passa o callback para atualizar a UI
        csv_str, kb, inventory, tags_by_id, tags_by_ip = tenable.get_as_csv_string_qualys_format(days_back=days_back, progress_callback=self.update_progress)

        # Retorna agora 5 valores para manter paridade com Qualys
        return csv_str, kb, inventory, tags_by_id, tags_by_ip

    def executing_qualys_logic(self, client_name, status_str, nvd_enabled):
        BASE_URL = os.getenv(f"QUALYS_URL_{client_name}") or os.getenv(f"QUALYS_BASE_URL_{client_name}")
        USERNAME = os.getenv(f"QUALYS_USER_{client_name}")
        PASSWORD = os.getenv(f"QUALYS_PASS_{client_name}")
        
        if not BASE_URL or not USERNAME or not PASSWORD:
            self.log("ERRO: Credenciais ausentes."); return None, {}, [], {}, {}

        session = requests.Session()
        session.auth = (USERNAME, PASSWORD)
        session.headers.update({'X-Requested-With': 'Python Script', 'Content-Type': 'application/x-www-form-urlencoded'})

        self.log(f">>> Buscando Inventário Qualys para {client_name} (Paginação Ativa)...")
        
        tags_by_id, tags_by_ip, all_ids, asset_inventory_local = {}, {}, [], []
        
        # --- NOVO BLOCO DE PAGINAÇÃO DE INVENTÁRIO (FIX LIMITE 1000) ---
        last_seen_id = 0
        page_count = 1
        
        while True:
            self.update_progress(0.1, f"Baixando Inventário - Página {page_count}...")
            
            # Payload solicitando 1000 assets a partir do último ID visto
            payload = {
                'action': 'list', 
                'details': 'All', 
                'show_tags': '1',
                'truncation_limit': '1000',  # Limite seguro por página
                'id_min': str(last_seen_id + 1)
            }
            
            try:
                r = session.post(f"{BASE_URL}/api/2.0/fo/asset/host/", data=payload, timeout=120)
                if r.status_code != 200: 
                    self.log(f"Erro API Qualys (Inventário): {r.status_code}")
                    break
                
                # Parsing XML
                try:
                    root = ET.fromstring(r.text)
                except ET.ParseError:
                    root = ET.fromstring(r.content)

                # Remove Namespaces
                for elem in root.iter():
                    if '}' in elem.tag:
                        elem.tag = elem.tag.split('}', 1)[1]

                hosts_in_batch = root.findall('.//HOST')
                if not hosts_in_batch:
                    self.log(f"   > Página {page_count}: Nenhum host retornado. Fim do inventário.")
                    break

                self.log(f"   > Página {page_count}: {len(hosts_in_batch)} hosts encontrados.")

                batch_max_id = 0
                
                for host in hosts_in_batch:
                    hid_node = host.find('ID')
                    if hid_node is None: continue
                    
                    hid = int(hid_node.text) # Converter para INT para achar o máximo
                    if hid > batch_max_id: batch_max_id = hid
                    
                    all_ids.append(str(hid))
                    
                    # IP Parsing
                    ip_node = host.find('IP')
                    if ip_node is None: ip_node = host.find('IP_ADDRESS')
                    ip = ip_node.text if ip_node is not None else "N/A"
                    
                    # Tags Parsing
                    tags = []
                    tags_node = host.find('TAGS')
                    if tags_node is not None:
                        # CORREÇÃO AQUI: "for tag_elem" (com espaço)
                        for tag_elem in tags_node.findall('TAG'):
                            name_node = tag_elem.find('NAME')
                            if name_node is not None and name_node.text: tags.append(name_node.text)
                        
                        if not tags: # Fallback option
                            for opt_elem in tags_node.findall('OPTION'):
                                name_node = opt_elem.find('NAME')
                                if name_node is not None and name_node.text: tags.append(name_node.text)
                    
                    tags_str = ", ".join(sorted(tags))
                    if str(hid) not in tags_by_id: tags_by_id[str(hid)] = tags_str
                    if ip and ip != "N/A": tags_by_ip[ip] = tags_str
                    
                    # Dados auxiliares
                    dns_node = host.find('DNS')
                    os_node = host.find('OS')
                    
                    asset_inventory_local.append({
                        'Host ID': str(hid), 
                        'IP Address': ip, 
                        'DNS Name': dns_node.text if dns_node is not None else "N/A", 
                        'Operating System': os_node.text if os_node is not None else "N/A", 
                        'Tags': tags_str
                    })

                # Verifica se precisamos continuar (se veio menos que o limite, acabou)
                if len(hosts_in_batch) < 1000:
                    break
                
                # Prepara próxima página
                last_seen_id = batch_max_id
                page_count += 1
                
            except Exception as e:
                self.log(f"Erro XML Parse na página {page_count}: {e}")
                break
        
        # --- FIM DA PAGINAÇÃO DO INVENTÁRIO ---

        if not all_ids: 
            self.log("Nenhum host Qualys encontrado após varredura."); return None, {}, [], {}, {}

        self.log(f"Total Final de Ativos: {len(all_ids)}. Iniciando download de detecções...")
        
        full_csv_lines = []
        unique_qids = set()
        
        # --- DEFINIÇÃO DA FUNÇÃO INTERNA DE BATCH (MANTIDA IGUAL) ---
        # --- FUNÇÃO DE BATCH COM PERSISTÊNCIA INFINITA ---
        def fetch_detections_batch(session_obj, id_list, batch_index, total_batches):
            url = f"{BASE_URL}/api/3.0/fo/asset/host/vm/detection/"
            ids_string = ",".join(id_list)
            
            payload = {
                'action': 'list', 'output_format': 'CSV', 'ids': ids_string,
                'truncation_limit': '0', 'status': status_str, 'show_qds': '1'
            }
            
            attempt = 1
            while True: # Loop infinito até o sucesso
                try:
                    # Timeout alto para garantir
                    r = session_obj.post(url, data=payload, timeout=300)
                    
                    if r.status_code == 200:
                        self.log(f"   [OK] Vulns Lote {batch_index}/{total_batches} (Tentativa {attempt})")
                        return r.text
                    
                    # Erros Fatais (Não adianta tentar de novo)
                    elif r.status_code in [400, 401, 403]:
                        self.log(f"   [ERRO FATAL] Lote {batch_index}: Código {r.status_code}. Abortando este lote.")
                        return None
                    
                    # Erros Transientes (Tenta de novo indefinidamente)
                    else:
                        # Calculamos um tempo de espera progressivo (Backoff)
                        # Ex: 15s, 30s, 45s... limitado a 2 minutos máximo
                        sleep_time = min(15 * attempt, 120) + random.randint(1, 10)
                        
                        error_type = "Rate Limit" if r.status_code == 429 else "Conflito" if r.status_code == 409 else f"Status {r.status_code}"
                        self.log(f"   [RETRY] Lote {batch_index}: {error_type}. Aguardando {sleep_time}s para tentativa {attempt + 1}...")
                        
                        time.sleep(sleep_time)
                        attempt += 1

                except Exception as e:
                    # Erro de Conexão/Rede -> Também tenta de novo
                    sleep_time = min(10 * attempt, 60)
                    self.log(f"   [ERRO REDE] Lote {batch_index}: {str(e)[:50]}... Tentando em {sleep_time}s (Tentativa {attempt + 1})")
                    time.sleep(sleep_time)
                    attempt += 1
        # -------------------------------------------------

        
        BATCH_SIZE = 250
        total_batches = (len(all_ids) + BATCH_SIZE - 1) // BATCH_SIZE
        completed_batches = 0
        
        future_to_batch = {}

        with ThreadPoolExecutor(max_workers=2) as executor:
            for i in range(0, len(all_ids), BATCH_SIZE):
                batch = all_ids[i:i+BATCH_SIZE]
                batch_num = (i // BATCH_SIZE) + 1
                
                # Iniciamos a thread e guardamos a referência no dicionário
                fut = executor.submit(fetch_detections_batch, session, batch, batch_num, total_batches)
                future_to_batch[fut] = batch_num

            # Processamento dos resultados conforme eles chegam
            for future in as_completed(future_to_batch):
                batch_num = future_to_batch[future] # Recupera qual lote é esse
                completed_batches += 1
                
                pct = 0.2 + (0.4 * (completed_batches / total_batches))
                
                try:
                    raw = future.result()
                    
                    if raw:
                        # SUCESSO
                        self.update_progress(pct, f"Baixando Detections ({completed_batches}/{total_batches})...")
                        lines = raw.strip().splitlines()
                        start = next((i for i, l in enumerate(lines) if '"QID"' in l.upper() or ",QID," in l.upper()), -1)
                        if start != -1:
                            full_csv_lines.extend(lines[start+1:] if full_csv_lines else lines[start:])
                            for l in lines[start:]: unique_qids.update(re.findall(r'\b\d{4,8}\b', l))
                    else:
                        # FALHA TOTAL DO LOTE (O None que estava sumindo)
                        self.log(f"CRÍTICO: O Lote {batch_num}/{total_batches} FALHOU completamente após retentativas. Dados perdidos.")
                        self.update_progress(pct, f"ERRO no Lote {batch_num}!")
                
                except Exception as e:
                    self.log(f"EXCEPTION CRÍTICA processando lote {batch_num}: {e}")
        
        csv_str = "\n".join(full_csv_lines)
        
        # --- LÓGICA DE CACHE PARA KNOWLEDGEBASE (MANTIDA) ---
        self.update_progress(0.65, "Verificando Cache de KnowledgeBase...")
        
        kb_cache = PersistentQualysKBCache()
        kb_details = {}
        missing_qids = []

        list_unique_qids = list(unique_qids)
        for qid in list_unique_qids:
            cached = kb_cache.get_qid(qid)
            if cached: kb_details[qid] = cached
            else: missing_qids.append(qid)

        self.log(f"KB Cache Hit: {len(kb_details)} | KB Cache Miss: {len(missing_qids)}")

        if missing_qids:
            KB_CHUNK_SIZE = 300 
            total_missing = len(missing_qids)
            
            def fetch_kb_chunk_with_retry(qids_chunk, attempt=1):
                chunk_details = {}
                try:
                    r = session.post(f"{BASE_URL}/api/2.0/fo/knowledge_base/vuln/", 
                                   data={'action': 'list', 'details': 'All', 'ids': ",".join(map(str, qids_chunk))}, 
                                   timeout=300)
                    
                    if r.status_code == 200:
                        try: root = ET.fromstring(r.text)
                        except: root = ET.fromstring(r.content)

                        for elem in root.iter():
                            if '}' in elem.tag: elem.tag = elem.tag.split('}', 1)[1]

                        for vuln in root.findall('.//VULN'):
                            qid_elem = vuln.find('QID')
                            if qid_elem is not None:
                                qid = qid_elem.text
                                cves = [c.find('ID').text for c in vuln.findall('.//CVE_LIST/CVE') if c.find('ID') is not None]
                                title = vuln.find('TITLE').text if vuln.find('TITLE') is not None else 'N/A'
                                category = vuln.find('CATEGORY').text if vuln.find('CATEGORY') is not None else 'N/A'
                                solution = vuln.find('SOLUTION').text if vuln.find('SOLUTION') is not None else 'Consulte o fabricante.'
                                
                                data = {'TITLE': title, 'CATEGORY': category, 'SOLUTION': solution, 'CVE_IDS': ",".join(cves)}
                                chunk_details[qid] = data
                        return chunk_details
                    elif r.status_code in [503, 429] and attempt <= 3:
                         time.sleep(10 * attempt); return fetch_kb_chunk_with_retry(qids_chunk, attempt + 1)
                except Exception:
                    if attempt <= 3: time.sleep(5); return fetch_kb_chunk_with_retry(qids_chunk, attempt + 1)
                return chunk_details

            processed_count = 0
            for i in range(0, total_missing, KB_CHUNK_SIZE):
                chunk = missing_qids[i:i+KB_CHUNK_SIZE]
                new_data = fetch_kb_chunk_with_retry(chunk)
                
                for qid, data in new_data.items():
                    kb_details[qid] = data
                    kb_cache.save_qid(qid, data)
                
                processed_count += len(chunk)
                pct = 0.65 + (0.25 * (processed_count / total_missing))
                self.update_progress(pct, f"Baixando KB: {processed_count}/{total_missing}...")

        session.close()
        
        return csv_str, kb_details, asset_inventory_local, tags_by_id, tags_by_ip
    
    def executar_logica(self, cliente_selecionado, status_str, segmento_selecionado, periodo_selecionado, usar_ia, start_date_str, usar_nvd):
        try:
            self.log("="*60); self.log(f"INICIANDO PROCESSO PARA: {cliente_selecionado}"); self.log("="*60)
            self.update_progress(0.05, "Iniciando Download de Dados...")
            
            csv_str = None
            kb_details = {}
            asset_inventory = []
            tags_by_id = {}
            tags_by_ip = {}
            
            is_tenable = os.getenv(f"TENABLE_ACCESS_KEY_{cliente_selecionado}")
            
            if is_tenable:
                self.log(">>> DETECTADO: TENABLE. Usando integração híbrida.")
                csv_str, kb_details, asset_inventory, tags_by_id, tags_by_ip = self.executing_tenable_logic(cliente_selecionado, usar_nvd, periodo_selecionado)
            else:
                self.log(">>> DETECTADO: QUALYS. Usando integração nativa.")
                csv_str, kb_details, asset_inventory, tags_by_id, tags_by_ip = self.executing_qualys_logic(cliente_selecionado, status_str, usar_nvd)

            if not csv_str:
                 self.log("Nenhum dado recebido."); self.after(0, lambda: self.finalizar_execucao(False)); return

            self.update_progress(0.7, "Download Concluído. Iniciando Processamento...")

            if NVD_API_KEYS and usar_nvd:
                self.log("Enriquecendo com NVD (Threat Intel)...")
                self.update_progress(0.75, "Iniciando Análise Threat Intel (NVD)...")
                
                nvd = NVDIntegration(NVD_API_KEYS)
                enricher = ThreatIntelligenceEnricher(nvd)
                qids_w_cve = [q for q, d in kb_details.items() if d.get('CVE_IDS')]
                total_enrich = len(qids_w_cve)
                completed_enrich = 0
                
                if total_enrich > 0:
                    with ThreadPoolExecutor(max_workers=min(len(NVD_API_KEYS)*10, 60)) as exc:
                        futures = {exc.submit(enricher.enrich_vulnerability, kb_details[q]): q for q in qids_w_cve}
                        for f in as_completed(futures):
                            completed_enrich += 1
                            if completed_enrich % 10 == 0: 
                                pct = 0.75 + (0.15 * (completed_enrich / total_enrich))
                                self.update_progress(pct, f"Analisando CVEs/Exploits ({completed_enrich}/{total_enrich})...")
                            try:
                                intel = f.result()
                                kb_details[futures[f]].update({
                                    'Threat_Max_CVSS': intel['max_cvss'], 
                                    'Threat_Exploit': "SIM" if intel['has_exploit'] else "NAO", 
                                    'Threat_KEV': "SIM" if intel['cisa_kev'] else "NAO", 
                                    'CVSS_Vector': intel['cvss_vector']
                                })
                            except: pass

            self.update_progress(0.9, "Gerando Relatórios e Excel...")
            
            source_type = "tenable" if is_tenable else "qualys"
            
            # Salva o CSV (Necessário para análise) e o Excel (Final)
            csv_path, excel_path = self.enrich_and_save_csv(
                csv_str, kb_details, cliente_selecionado, 
                tags_by_id, tags_by_ip, periodo_selecionado, 
                asset_inventory, source_type=source_type
            )
            
            if csv_path:
                # O CSV ainda existe aqui e é usado para análise
                summary = self.analyze_data_for_charts(csv_path, periodo_selecionado)
                if summary:
                    self.log("Gerando Dashboard Executivo...")
                    # Passa o diretório do CSV para salvar os gráficos lá também
                    chart_gen = ReportChartGenerator(os.path.dirname(csv_path))
                    chart_path = chart_gen.create_executive_summary_chart(summary)
                    
                    ai_text = "Análise ignorada."
                    if usar_ia:
                        self.log("Consultando IA...")
                        self.update_progress(0.95, "Consultando Inteligência Artificial...")
                        ai_text, model = self.perform_ai_analysis(summary, segmento_selecionado, periodo_selecionado, start_date_str)
                        self.log(f"IA gerada por: {model}")

                    self.generate_word_report_v2(cliente_selecionado, summary, ai_text, chart_path, os.path.dirname(csv_path), periodo_selecionado)
                    
                    try:
                        # Carrega dados para a Web antes de deletar o arquivo
                        df = pd.read_csv(csv_path, on_bad_lines='skip', low_memory=False)
                        df = self.filtrar_dataframe_por_data(df, periodo_selecionado)
                        shared_data.dataframe = df
                        shared_data.client_name = cliente_selecionado
                        shared_data.updated_at = datetime.now()
                        self.log("Dashboard Web Atualizado.")
                    except: pass

                    # --- EXCLUSÃO DO CSV TEMPORÁRIO ---
                    try:
                        os.remove(csv_path)
                        self.log(f"Arquivo temporário CSV excluído: {csv_path}")
                    except Exception as e:
                        self.log(f"Aviso: Não foi possível excluir o CSV temporário: {e}")
                    # ----------------------------------
                    
                    self.after(0, lambda: self.finalizar_execucao(True, "Processo Finalizado!"))
            else:
                self.after(0, lambda: self.finalizar_execucao(False, "Falha ao salvar arquivos."))

        except Exception as e:
            self.log(f"ERRO FATAL: {e}"); traceback.print_exc(); self.after(0, lambda: self.finalizar_execucao(False))


    def enrich_and_save_csv(self, all_csv_data, kb_details, cliente_name, tags_by_id, tags_by_ip, periodo_str, asset_inventory, source_type="qualys"):
        try: csv.field_size_limit(sys.maxsize)
        except OverflowError: csv.field_size_limit(2147483647) 

        self.log(f"Cruzando dados e gerando Excel Formatado ({source_type.upper()})...")
        csv_file = StringIO(all_csv_data)
        reader = csv.reader(csv_file)
        
        try:
            # 1. Identificar Cabeçalho Original
            header_found = False
            source_header = []
            for line in reader:
                if not line: continue
                line_str = ",".join(line).upper()
                if "IP" in line_str and ("QID" in line_str or "PLUGIN" in line_str):
                    source_header = line
                    header_found = True
                    break
            
            if not header_found: return None, None
            
            cleaned_header = [h.strip().strip('"').upper() for h in source_header]
            
            # 2. Mapeamento de índices
            def get_idx(name_list):
                for name in name_list:
                    try: return cleaned_header.index(name)
                    except: continue
                return -1

            qid_idx = get_idx(['QID'])
            tenable_id_idx = get_idx(['TENABLE_PLUGIN_ID', 'PLUGIN_ID'])
            vpr_idx = get_idx(['VPR_SCORE', 'VPR'])
            cvss3_idx = get_idx(['CVSS_V3_BASE', 'CVSS3_BASE_SCORE'])
            risk_idx = get_idx(['RISK_FACTOR', 'RISK'])
            host_id_idx = get_idx(["HOST ID", "QUALYS HOST ID"])
            ip_col_idx = get_idx(["IP", "IP ADDRESS", "IPADDRESS"])
            dns_idx = get_idx(["DNS", "DNS NAME"])
            os_idx = get_idx(["OS", "OPERATING SYSTEM"])
            sev_idx = get_idx(["SEVERITY", "SEV"])
            status_idx = get_idx(["STATUS", "STATE"])
            first_found_idx = get_idx(["FIRST FOUND", "FIRST_FOUND"])
            last_found_idx = get_idx(["LAST DETECTED", "LAST_FOUND"])

            # 3. Definição do Cabeçalho Final
            if source_type == 'tenable':
                final_header = [
                    "Plugin ID", "VPR Score", "Risk Factor", "CVSS v3 Base", "Severity", 
                    "IP Address", "DNS Name", "OS", "Status", 
                    "First Found", "Last Detected", 
                    "Tags", "Title", "Solution", "CVEs", 
                    "Threat_Exploit", "Threat_KEV", "Threat_Max_CVSS", "CVSS_Vector"
                ]
            else:
                final_header = source_header + ["ASSET_TAGS", "Vuln_TITLE", "Vuln_CATEGORY", "Vuln_SOLUTION", "CVEs", "Threat_Exploit", "Threat_KEV", "Threat_Max_CVSS", "CVSS_Vector"]
            
            # --- ALTERAÇÃO: ESTRUTURA DE PASTAS (Ano/Mês) ---
            agora = datetime.now()
            ano_atual = str(agora.year)
            
            # Mapeamento manual para garantir nomes em Português independente do SO
            mapa_meses = {
                1: "01-Janeiro", 2: "02-Fevereiro", 3: "03-Março", 4: "04-Abril",
                5: "05-Maio", 6: "06-Junho", 7: "07-Julho", 8: "08-Agosto",
                9: "09-Setembro", 10: "10-Outubro", 11: "11-Novembro", 12: "12-Dezembro"
            }
            mes_pasta = mapa_meses[agora.month]
            
            # Cria caminho: Clientes \ Cliente \ 2024 \ 01-Janeiro
            output_folder = os.path.join("Clientes", cliente_name, ano_atual, mes_pasta)
            os.makedirs(output_folder, exist_ok=True)
            
            # Formato do arquivo: 22JAN26-11h17m
            timestamp_formatted = agora.strftime("%d%b%y-%Hh%Mm").upper()
            
            base_filename = f"{cliente_name}-{timestamp_formatted}"
            
            final_csv_file = os.path.join(output_folder, f"{base_filename}.csv")
            final_excel_file = os.path.join(output_folder, f"{base_filename}.xlsx")
            # ------------------------------------------------

            # 4. Escrita do CSV Final
            with open(final_csv_file, 'w', newline='', encoding='utf-8') as outfile:
                writer = csv.writer(outfile)
                writer.writerow(final_header)
                
                for row in reader:
                    if source_type == 'tenable' and tenable_id_idx != -1:
                        main_id = str(row[tenable_id_idx]).strip()
                    elif qid_idx != -1 and len(row) > qid_idx:
                        main_id = str(row[qid_idx]).strip()
                    else: continue 

                    tags_str = ""
                    if host_id_idx != -1 and len(row) > host_id_idx:
                        h_id = row[host_id_idx]
                        if h_id in tags_by_id: tags_str = tags_by_id[h_id]
                    
                    if not tags_str and ip_col_idx != -1 and len(row) > ip_col_idx:
                        ip_val = row[ip_col_idx].strip().strip('"')
                        if ip_val in tags_by_ip: tags_str = tags_by_ip[ip_val]
                        elif ip_val.strip() in tags_by_ip: tags_str = tags_by_ip[ip_val.strip()]
                    
                    details = kb_details.get(main_id, {})
                    title = details.get('TITLE', 'N/A')
                    category = details.get('CATEGORY', 'N/A')
                    solution = details.get('SOLUTION', 'N/A')
                    cves = details.get('CVE_IDS', '')
                    exploit = details.get('Threat_Exploit', '')
                    kev = details.get('Threat_KEV', '')
                    cvss_threat = details.get('Threat_Max_CVSS', '')
                    vector = details.get('CVSS_Vector', '')

                    if source_type == 'tenable':
                        def safe_get(idx): return row[idx] if idx != -1 and idx < len(row) else ""
                        new_row = [
                            main_id, safe_get(vpr_idx), safe_get(risk_idx), safe_get(cvss3_idx), safe_get(sev_idx), 
                            safe_get(ip_col_idx), safe_get(dns_idx), safe_get(os_idx), safe_get(status_idx), 
                            safe_get(first_found_idx), safe_get(last_found_idx), 
                            tags_str, title, solution, cves, 
                            exploit, kev, cvss_threat, vector
                        ]
                        writer.writerow(new_row)
                    else:
                        writer.writerow(row + [tags_str, title, category, solution, cves, exploit, kev, cvss_threat, vector])

            # 5. Geração do Excel
            df = pd.read_csv(final_csv_file, on_bad_lines='skip', low_memory=False)
            df_filtrado = self.filtrar_dataframe_por_data(df, periodo_str)
            self.log(f"Linhas após filtro de data: {len(df_filtrado)}")

            df_inventory = pd.DataFrame(asset_inventory)
            if not df_inventory.empty and 'Tags' in df_inventory.columns:
                unique_tags = set()
                for t_str in df_inventory['Tags'].dropna():
                    if t_str: unique_tags.update([p.strip() for p in t_str.split(',')])
                for tag in sorted(list(unique_tags)):
                    df_inventory[tag] = df_inventory['Tags'].apply(lambda x: tag if x and tag in str(x).split(',') else "")

            def add_formatted_sheet(writer, df_to_write, sheet_name, tab_color, table_style='Table Style Medium 2'):
                if df_to_write.empty: return
                df_to_write = df_to_write.copy() 
                for col in df_to_write.columns:
                    if pd.api.types.is_datetime64_any_dtype(df_to_write[col]):
                        try: df_to_write[col] = df_to_write[col].dt.tz_localize(None)
                        except: pass
                df_to_write.columns = df_to_write.columns.astype(str)
                df_to_write.to_excel(writer, sheet_name=sheet_name, index=False, startrow=1, header=False)
                worksheet = writer.sheets[sheet_name]
                worksheet.set_tab_color(tab_color)
                (max_row, max_col) = df_to_write.shape
                column_settings = [{'header': col} for col in df_to_write.columns]
                worksheet.add_table(0, 0, max_row, max_col - 1, {'columns': column_settings, 'style': table_style, 'name': re.sub(r'\W+', '_', sheet_name)})
                worksheet.set_column(0, max_col - 1, 18)

            cols = df_filtrado.columns
            col_status_name = next((c for c in cols if c.upper() in ["STATUS", "STATE"]), "Status")
            col_title_name = next((c for c in cols if c.upper() in ["TITLE", "VULN_TITLE"]), "Title")
            col_exploit_name = next((c for c in cols if "THREAT_EXPLOIT" in c.upper()), "Threat_Exploit")

            with pd.ExcelWriter(final_excel_file, engine='xlsxwriter') as writer:
                if not df_inventory.empty:
                    add_formatted_sheet(writer, df_inventory, 'INVENTÁRIO', 'blue', 'Table Style Medium 9')
                if col_status_name in df_filtrado.columns:
                    df_fixed = df_filtrado[df_filtrado[col_status_name].astype(str).str.upper().str.contains('FIXED')]
                    add_formatted_sheet(writer, df_fixed, 'FIXED', '#008000', 'Table Style Medium 7')
                if col_title_name in df_filtrado.columns:
                    mask_eol = df_filtrado[col_title_name].astype(str).str.contains(r'EOL|Obsolete|End of Life|Unsupported', case=False, regex=True)
                    df_eol = df_filtrado[mask_eol]
                    add_formatted_sheet(writer, df_eol, 'EOL', '#800080', 'Table Style Medium 12')
                if col_exploit_name in df_filtrado.columns:
                      df_exploit = df_filtrado[df_filtrado[col_exploit_name] == "SIM"]
                      if not df_exploit.empty:
                          add_formatted_sheet(writer, df_exploit, 'CRITICAL_EXPLOIT', '#C0392B', 'Table Style Medium 3')
                if col_status_name in df_filtrado.columns:
                    df_active = df_filtrado[~df_filtrado[col_status_name].astype(str).str.upper().str.contains('FIXED')]
                    add_formatted_sheet(writer, df_active, 'Todas ATIVAS', '#FF0000', 'Table Style Medium 3')

            return final_csv_file, final_excel_file

        except Exception as e:
            self.log(f"Erro na geração do Excel: {e}")
            traceback.print_exc()
            return None, None

    def filtrar_dataframe_por_data(self, df, periodo_str):
        if not periodo_str or "Todo" in periodo_str: return df
        try:
            dias = int(periodo_str.split()[0])
            
            # Tenta achar a coluna de Last Detected
            cols_possiveis = [c for c in df.columns if 'LAST' in c.upper() and ('FOUND' in c.upper() or 'DETECTED' in c.upper())]
            if not cols_possiveis: return df
            
            col_last = cols_possiveis[0]
            
            # Converte forçando erros a virarem NaT, mas assumindo dia primeiro se ambíguo
            df[col_last] = pd.to_datetime(df[col_last], errors='coerce', dayfirst=True, utc=True)
            
            # Se a conversão falhou em tudo (tudo NaT), retorna o df original para não zerar o dashboard
            if df[col_last].notna().sum() == 0:
                self.log("Aviso: Não foi possível filtrar por data (formato desconhecido). Exibindo tudo.")
                return df

            cutoff = pd.Timestamp.now(timezone.utc) - pd.Timedelta(days=dias)
            
            # Filtra onde a data é maior que o cutoff OU onde a data é NaT (para não sumir com erros de parse)
            return df[(df[col_last] >= cutoff) | (df[col_last].isna())]
        except Exception as e:
            self.log(f"Erro no filtro de data: {e}")
            return df


    def analyze_data_for_charts(self, file_path, periodo_str):
        try:
            df = pd.read_csv(file_path, on_bad_lines='skip', low_memory=False)
            df = self.filtrar_dataframe_por_data(df, periodo_str)
            if df.empty: return None
        except: return None
        
        col_sev = next((c for c in df.columns if 'SEV' in c.upper()), None)
        col_title = next((c for c in df.columns if 'TITLE' in c.upper()), None)
        col_status = next((c for c in df.columns if 'STATUS' in c.upper()), None) 
        col_ff = next((c for c in df.columns if 'FIRST' in c.upper()), None)
        col_fixed = next((c for c in df.columns if 'FIXED' in c.upper() and 'DATE' in c.upper()), None)
        
        if col_sev:
            df['Severity'] = pd.to_numeric(df[col_sev], errors='coerce').fillna(0)
        else:
            df['Severity'] = 0
        
        df['Age_Days'] = 0
        if col_ff:
            try:
                df['First Found'] = pd.to_datetime(df[col_ff], errors='coerce', utc=True)
                df['Age_Days'] = (pd.Timestamp.now(timezone.utc) - df['First Found']).dt.days.fillna(0)
            except: pass
            
        if col_fixed:
            try:
                df['Fixed_Date'] = pd.to_datetime(df[col_fixed], errors='coerce', utc=True)
            except: pass

        df['Risk_Score'] = df['Severity'] * 10
        if col_title:
            crit = df[col_title].astype(str).str.upper().str.contains(r'RCE|SQL|RANSOMWARE', regex=True)
            df.loc[crit, 'Risk_Score'] *= 2

        stats = {'Active': 0, 'Fixed': 0, 'Re-Opened': 0, 'New': 0}
        if col_status:
            s = df[col_status].astype(str).str.upper()
            stats['Fixed'] = s[s.str.contains('FIXED', na=False)].count()
            stats['Active'] = s[s.str.contains('ACTIVE', na=False)].count()
            stats['New'] = s[s.str.contains('NEW', na=False)].count()
            stats['Re-Opened'] = s[s.str.contains('RE-OPENED', na=False)].count()

        severity_counts = df['Severity'].value_counts().to_dict()

        aging = {
            '0-30': len(df[df['Age_Days'] <= 30]),
            '31-60': len(df[(df['Age_Days'] > 30) & (df['Age_Days'] <= 60)]),
            '61-90': len(df[(df['Age_Days'] > 60) & (df['Age_Days'] <= 90)]),
            '90+': len(df[df['Age_Days'] > 90])
        }

        avg_sev = df['Severity'].mean()
        total_vulns = len(df)
        ratio_crit = 0
        if total_vulns > 0:
            if col_status:
                crit_active = len(df[(df['Severity'] >= 4) & (~df[col_status].astype(str).str.upper().str.contains('FIXED', na=False))])
            else:
                crit_active = len(df[df['Severity'] >= 4])
            ratio_crit = crit_active / total_vulns
        
        risk_norm = (avg_sev / 5 * 40) + (ratio_crit * 60)
        risk_norm = min(risk_norm * 1.8, 100) 

        top_remediation = []
        try:
            grp_cols = []
            if 'QID' in df.columns: grp_cols.append('QID')
            if col_title: grp_cols.append(col_title)
            if not grp_cols: grp_cols = [df.columns[0]] 

            col_ip = next((c for c in df.columns if 'IP' in c.upper()), None)
            if not col_ip: col_ip = df.columns[0] 
            
            df_active = df
            if col_status:
                df_active = df[~df[col_status].astype(str).str.upper().str.contains('FIXED', na=False)]

            if not df_active.empty:
                remed = df_active[df_active['Severity'] >= 3].groupby(grp_cols).agg(
                    Hosts_Count=(col_ip, 'nunique'),
                    Max_Sev=('Severity', 'max'),
                    Risk_Sum=('Risk_Score', 'sum')
                ).reset_index().sort_values('Risk_Sum', ascending=False).head(20)
                
                if col_title:
                    remed.rename(columns={col_title: 'Vuln_TITLE'}, inplace=True)
                else:
                    remed['Vuln_TITLE'] = "N/A"

                top_remediation = remed.to_dict('records')
        except: pass

        mttr = {}
        if col_fixed and not df.empty:
            try:
                df_fixed = df[df['Fixed_Date'].notnull()].copy()
                if not df_fixed.empty:
                    df_fixed['Time_To_Fix'] = (df_fixed['Fixed_Date'] - df_fixed['First Found']).dt.days
                    mttr = df_fixed.groupby('Severity')['Time_To_Fix'].mean().round(1).to_dict()
            except: pass

        return {
            'status_counts': stats,
            'severity_counts': severity_counts,
            'aging_buckets': aging,
            'risk_score_norm': risk_norm,
            'top_remediation': top_remediation,
            'avg_time_to_fix': mttr,
            'smart_insights': {'old_crit_count': len(df[(df['Age_Days'] > 90) & (df['Severity'] >= 4)])}
        }

    def generate_word_report_v2(self, client, summary, ai_text, chart_path, out_dir, periodo):
        try:
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            doc.add_paragraph("\n" * 4)
            title_head = doc.add_heading(f'Relatório Executivo de Segurança', 0)
            title_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_paragraph(f"Análise de Vulnerabilidades & Risco Cibernético")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(16)
            subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128) 
            
            doc.add_paragraph("\n" * 2)
            table_cover = doc.add_table(rows=3, cols=1)
            table_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row0 = table_cover.rows[0].cells[0]
            p0 = row0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run0 = p0.add_run(f"CLIENTE: {client.upper()}")
            run0.bold = True
            run0.font.size = Pt(14)
            
            row1 = table_cover.rows[1].cells[0]
            p1 = row1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.add_run(f"Período Analisado: {periodo}")
            
            row2 = table_cover.rows[2].cells[0]
            p2 = row2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run(f"Gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
            
            doc.add_page_break()
            
            doc.add_heading('1. Panorama Visual (Dashboard)', 1)
            if chart_path and os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.2))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph("\n")

            doc.add_heading('2. Indicadores Críticos (KPIs)', 1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Medium Grid 1 Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Vulnerabilidades Ativas'
            hdr_cells[1].text = 'Vulnerabilidades Corrigidas'
            hdr_cells[2].text = 'Legado Crítico (>90d)'
            
            row_cells = table.add_row().cells
            
            def format_cell_big(cell, text, color=None):
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(text))
                run.bold = True
                run.font.size = Pt(24)
                if color: run.font.color.rgb = color
            
            format_cell_big(row_cells[0], summary['status_counts']['Active'], RGBColor(200, 0, 0)) 
            format_cell_big(row_cells[1], summary['status_counts']['Fixed'], RGBColor(0, 128, 0)) 
            format_cell_big(row_cells[2], summary['smart_insights']['old_crit_count'], RGBColor(255, 140, 0)) 
            
            doc.add_paragraph("\n")
            
            doc.add_heading('3. Análise Estratégica & Priorização (AI)', 1)
            intro = doc.add_paragraph()
            intro_run = intro.add_run("A análise abaixo foi gerada por Inteligência Artificial Generativa baseada no contexto do segmento e nos dados técnicos coletados.")
            intro_run.italic = True
            intro_run.font.size = Pt(9)
            intro_run.font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_paragraph("-" * 90)
            self._apply_markdown_formatting(doc, ai_text)
            
            # --- ALTERAÇÃO: CAMINHO E NOME DO ARQUIVO DOC ---
            agora = datetime.now()
            ano_atual = str(agora.year)
            mapa_meses = {
                1: "01-Janeiro", 2: "02-Fevereiro", 3: "03-Março", 4: "04-Abril",
                5: "05-Maio", 6: "06-Junho", 7: "07-Julho", 8: "08-Agosto",
                9: "09-Setembro", 10: "10-Outubro", 11: "11-Novembro", 12: "12-Dezembro"
            }
            mes_pasta = mapa_meses[agora.month]
            
            # Cria/Garante caminho: Clientes \ Cliente \ 2024 \ 01-Janeiro
            output_folder = os.path.join("Clientes", client, ano_atual, mes_pasta)
            os.makedirs(output_folder, exist_ok=True)

            timestamp_formatted = agora.strftime("%d%b%y-%Hh%Mm").upper()
            
            fname = os.path.join(output_folder, f"{client}-{timestamp_formatted}.docx")
            # ------------------------------------------------
            
            doc.save(fname)
            return fname
        except Exception as e:
            traceback.print_exc()
            return f"Erro Word: {e}"

    def _apply_markdown_formatting(self, doc, text):
        for line in text.split('\n'):
            line = line.strip()
            if not line: continue
            if line.startswith('### '):
                h = doc.add_heading(line.replace('### ', ''), level=3)
                h.runs[0].font.color.rgb = RGBColor(50, 50, 50)
            elif line.startswith('## '):
                h = doc.add_heading(line.replace('## ', ''), level=2)
                h.runs[0].font.color.rgb = RGBColor(30, 60, 100)
            elif line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                clean_text = line[2:]
                self._parse_bold_runs(p, clean_text)
            else:
                p = doc.add_paragraph()
                self._parse_bold_runs(p, line)

    def _parse_bold_runs(self, paragraph, text):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part.replace('**', ''))
                run.bold = True
            else:
                paragraph.add_run(part)

if __name__ == "__main__":
    app = ModernQualysApp()
    app.mainloop()
